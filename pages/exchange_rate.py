import streamlit as st
import os
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import yfinance as yf
import requests
import time
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from io import BytesIO
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import plotly.graph_objects as go
from urllib.parse import unquote
import base64

# --- [1. 페이지 기본 설정] ---
st.set_page_config(page_title="Trade Master 2026", layout="wide", page_icon="🚢")

# --- [2. 환경 변수 및 OpenAI 설정] ---
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
kotra_api_key = os.getenv("KOTRA_API_KEY")

if not api_key:
    st.warning("⚠️ Open_api_key 환경 변수가 설정되지 않았습니다. AI 기능이 제한될 수 있습니다.")
    client = None
else:
    try:
        client = OpenAI(api_key=api_key, timeout=60.0)
    except Exception as e:
        st.error(f"OpenAI 클라이언트 초기화 오류: {e}")
        client = None

plt.rcParams['font.family'] = 'Pretendard'
plt.rcParams['axes.unicode_minus'] = False

# --- [3. UI 디자인 및 스타일링] ---
st.markdown("""
    <style>
    /* Streamlit 기본 멀티페이지 네비게이션 제거 */
    [data-testid="stSidebarNav"] {
        display: none;
    }
    
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
    html, body, [class*="css"], .stApp {
        font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, system-ui, Roboto, sans-serif !important;
    }
    .stApp { background-color: #ffffff; color: #31333f; }
    section[data-testid="stSidebar"] { background-color: #ffffff !important; border-right: 1px solid #e5e7eb; }

    .stButton>button {
        width: 100%;
        border-radius: 8px;
        height: 3em;
        background-color: #051161;
        color: white;
        border: none;
        transition: 0.3s;
        font-weight: 600;
    }
    .stButton>button:hover {
        background-color: rgba(5,17,97,0.85);
        box-shadow: 0 4px 12px rgba(5,17,97,0.3);
    }

    /* --- sidebar logo styles --- */
    .logo-box{
        background: rgba(255,255,255,0.6);
        border: 1px solid #e5e7eb;
        border-radius: 14px;
        padding: 14px 12px;
        margin-bottom: 10px;
        text-align:center;
    }
    .logo-img{
        max-width: 150px;
        width: 100%;
        height: auto;
        display:block;
        margin: 0 auto;
    }
    .small-muted{
        color:#64748b;
        font-size: 0.85rem;
        font-weight: 600;
        letter-spacing: 0.2px;
    }
    </style>
    """, unsafe_allow_html=True)

# --- [4. 데이터 및 세션 상태 초기화] ---
if 'exchange_rates' not in st.session_state:
    st.session_state['exchange_rates'] = {"USD": 1440.70, "JPY": 935.94, "EUR": 1717.31, "CNY": 207.38}
if 'use_realtime' not in st.session_state:
    st.session_state['use_realtime'] = False

KCCI_FREIGHT_RATES = {
    "부산 -> LA": {"ocean_per_cbm": 145, "transit_days": 14},
    "인천 -> 상하이": {"ocean_per_cbm": 55, "transit_days": 2},
    "부산 -> 함부르크": {"ocean_per_cbm": 230, "transit_days": 35},
    "인천 -> 호치민": {"ocean_per_cbm": 78, "transit_days": 5}
}

# --- [5. 주요 기능 함수] ---

def get_realtime_exchange_rates():
    tickers = {"USD": "USDKRW=X", "JPY": "JPYKRW=X", "EUR": "EURKRW=X", "CNY": "CNYKRW=X"}
    updated_rates = {}
    try:
        for code, ticker in tickers.items():
            data = yf.download(ticker, period="2d", interval="1d", progress=False)
            if not data.empty:
                val = data['Close'].iloc[-1]
                updated_rates[code] = float(val) * 100 if code == "JPY" else float(val)
            else:
                updated_rates[code] = st.session_state['exchange_rates'][code]
        return updated_rates
    except Exception as e:
        st.error(f"환율 데이터를 가져오는 중 오류 발생: {e}")
        return st.session_state['exchange_rates']

@st.cache_data(ttl=3600)
def fetch_kotra_news(country_name):
    endpoint = "https://apis.data.go.kr/B410001/kotra_overseasMarketNews/ovseaMrktNews"
    decoded_key = unquote(kotra_api_key) if kotra_api_key else ""
    params = {'serviceKey': decoded_key, 'numOfRows': '5', 'pageNo': '1', 'search1': country_name}
    try:
        response = requests.get(endpoint, params=params, timeout=10)
        if response.status_code == 200:
            root = ET.fromstring(response.content)
            news_list = []
            for item in root.findall('.//item'):
                title = item.findtext('newsTitl')
                abst = item.findtext('newsAbst')
                clean_summary = BeautifulSoup(abst, "html.parser").get_text() if abst else ""
                news_list.append({
                    'title': title,
                    'url': f"https://dream.kotra.or.kr/dream/cms/news/actionOvseaMrktNewsDetail.do?SITE_NO=3&MENU_ID=180&CONTENTS_NO=1&bbsGbn=243&bbsSn={item.findtext('bbsSn')}",
                    'date': item.findtext('regDt'),
                    'summary': clean_summary[:100] + "..." if len(clean_summary) > 100 else clean_summary
                })
            return news_list
        return []
    except Exception:
        return []

@st.cache_data(ttl=3600)
def get_currency_history(ticker_symbol, base_val, multiplier, use_realtime, current_date):
    if use_realtime:
        try:
            data = yf.download(ticker_symbol, period="1mo", interval="1d", progress=False)
            if not data.empty:
                df = data.copy().reset_index()
                df.columns = ["날짜", "Open", "High", "Low", "Close", "Adj Close", "Volume"]
                for col in ["Open", "High", "Low", "Close"]:
                    df[col] = df[col] * multiplier
                df['날짜'] = pd.to_datetime(df['날짜']).dt.date
                df['환율'] = df['Close']
                return df.sort_values(by="날짜")
        except Exception:
            pass

    np.random.seed(abs(hash(ticker_symbol)) % (10**8))
    dates = pd.date_range(end=current_date, periods=30)
    close_vals = base_val + np.cumsum(np.random.randn(30) * (base_val * 0.005))
    df = pd.DataFrame({
        "날짜": dates.date,
        "Close": close_vals,
        "Open": close_vals * (1 + np.random.randn(30) * 0.002),
        "High": close_vals * (1 + np.random.rand(30) * 0.005),
        "Low": close_vals * (1 - np.random.rand(30) * 0.005),
        "환율": close_vals
    })
    return df

def calculate_estimated_cost(base_price, term, transport, insurance, payment_text, fta_type):
    total = base_price
    freight_bearing_terms = ["CFR", "CIF", "CPT", "CIP", "DAT", "DAP", "DPU", "DDP"]
    if term in freight_bearing_terms:
        freight_rate = 0.12 if transport == "항공(AIR)" else 0.04
        total += base_price * freight_rate

    ins_rates = {"ICC(A)": 0.008, "ICC(B)": 0.005, "ICC(C)": 0.003, "선택 안함": 0}
    if term in ["CIF", "CIP", "DDP"] or insurance != "선택 안함":
        total += base_price * ins_rates.get(insurance, 0)

    pay_fee_rate = 0.0
    upper_pay = payment_text.upper()
    if "L/C" in upper_pay:
        pay_fee_rate = 0.012
    elif "D/P" in upper_pay or "D/A" in upper_pay:
        pay_fee_rate = 0.005
    elif "T/T" in upper_pay:
        pay_fee_rate = 0.001
    total += base_price * pay_fee_rate

    fta_rates = {
        "협정 미적용 (기본세율 8%)": 0.08,
        "한-미 FTA (KOR-USA)": 0.00,
        "한-EU FTA (KOR-EU)": 0.00,
        "한-중 FTA (KOR-CHINA)": 0.04,
        "한-아세안 FTA (KOR-ASEAN)": 0.02,
        "한-베트남 FTA": 0.00,
        "한-인도 CEPA": 0.05,
        "한-영 FTA (KOR-UK)": 0.00,
        "RCEP (역내포괄적경제동반자협정)": 0.03,
        "한-호주 FTA": 0.00,
        "한-캐나다 FTA": 0.00
    }
    if "DDP" in term:
        total += base_price * fta_rates.get(fta_type, 0.08)

    return total

def calculate_quote(route, cbm, shipment_mode):
    route_data = KCCI_FREIGHT_RATES.get(route, KCCI_FREIGHT_RATES["부산 -> LA"])
    ocean_usd_per_cbm = route_data["ocean_per_cbm"]
    ex_rate = 1450.0
    if shipment_mode == "LCL (소량 화물)":
        ocean_cost_usd = cbm * (ocean_usd_per_cbm / 25 * 1.2)
    else:
        ocean_cost_usd = cbm * ocean_usd_per_cbm
    ocean_cost_krw = ocean_cost_usd * ex_rate
    return {"usd": round(ocean_cost_usd, 2), "krw": int(ocean_cost_krw), "days": route_data["transit_days"]}

def get_ai_exchange_insight(df_rates):
    """
    첫 번째 환율 표 데이터를 기반으로 AI 인사이트를 생성합니다.
    """
    if client is None:
        return "⚠️ OpenAI API 키가 설정되지 않아 인사이트를 생성할 수 없습니다."

    # 데이터프레임의 내용을 AI가 읽기 편한 텍스트로 변환
    rates_summary = ""
    for _, row in df_rates.iterrows():
        rates_summary += f"- {row['통화명']}: 매매기준율 {row['매매기준율']:.2f}원, 변동성 {row['시장 변동성 (%)']:.2f}%\n"

    prompt = f"""
    당신은 글로벌 무역 금융 전문가입니다. 아래 환율 데이터를 바탕으로 무역 종사자를 위한 핵심 인사이트를 3줄로 요약해 주세요.
    특히 어떤 통화가 강세/약세인지, 송금 시 주의할 점은 무엇인지 포함해 주세요.

    [실시간 환율 데이터]
    {rates_summary}
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ AI 분석 중 오류 발생: {e}"

def draw_candlestick_chart(df, label):
    fig = go.Figure(data=[go.Candlestick(
        x=df['날짜'],
        open=df['Open'],
        high=df['High'],
        low=df['Low'],
        close=df['Close'],
        increasing_line_color='#ef5350',
        decreasing_line_color='#26a69a',
        name=label
    )])
    fig.update_layout(
        title=dict(text=f"<b>{label} 분석</b>", font=dict(family='Pretendard', size=18, color='#1e293b')),
        template='plotly_white',
        margin=dict(l=40, r=40, t=60, b=40),
        height=450,
        xaxis_rangeslider_visible=False,
        font=dict(family='Pretendard'),
        yaxis=dict(tickformat=',.2f', title="환율 (KRW)", gridcolor='#f1f5f9'),
        xaxis=dict(type='date', gridcolor='#f1f5f9')
    )
    return fig

# --- [6. 서류 생성 관련 함수] ---

def create_common_table(doc, data):
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = f"Exporter (Seller):\n{data['shipper']}"
    table.rows[0].cells[1].text = f"Ref No. & Date:\n{data['inv_no_date']}"
    table.rows[1].cells[0].text = f"Consignee (Buyer):\n{data['consignee']}"
    table.rows[1].cells[1].text = f"L/C Issuing Bank:\n{data.get('lc_bank', 'N/A')}"
    table.rows[2].cells[0].text = f"Notify Party:\n{data.get('notify', 'Same as Consignee')}"
    table.rows[2].cells[1].text = f"Port of Loading / Discharge:\n{data['from_port']} / {data['to_port']}"
    table.rows[3].cells[0].text = f"Vessel / Flight No:\n{data['vessel']}"
    table.rows[3].cells[1].text = f"Terms of Delivery & Payment:\n{data['terms']} / {data['pay']}"
    return table

def create_item_table(doc, data):
    doc.add_paragraph("\n[ITEM LIST / DESCRIPTION]")
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "No.", "Description of Goods", "Quantity", "Unit Price (USD)", "Amount (USD)"
    row = table.rows[1].cells
    row[0].text = "1"
    row[1].text = data['description']
    row[2].text = str(data['qty'])
    row[3].text = str(data['unit_price'])
    row[4].text = str(data['amount'])
    return table

def create_sales_contract_docx(data):
    doc = Document()
    doc.add_heading('SALES CONTRACT', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    create_common_table(doc, data)
    create_item_table(doc, data)
    doc.add_paragraph("\n[TERMS & CONDITIONS]")
    doc.add_paragraph("1. Definitions and Scope: ... (생략)")
    doc.add_paragraph("5. Governing Law and Arbitration: This contract shall be governed by the laws of the Republic of Korea. Arbitration in Seoul under ICC rules.")
    return doc

def create_commercial_invoice_docx(data):
    doc = Document()
    doc.add_heading('COMMERCIAL INVOICE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    create_common_table(doc, data)
    create_item_table(doc, data)
    doc.add_paragraph("\n[TERMS & CONDITIONS]")
    doc.add_paragraph("9. Compliance: ... (생략)\n10. Governing Law: Republic of Korea.")
    return doc

def create_packing_list_docx(data):
    doc = Document()
    doc.add_heading('PACKING LIST', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    create_common_table(doc, data)
    create_item_table(doc, data)
    doc.add_paragraph("\n[PACKING & SHIPPING DETAILS]")
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Total Net Weight (N.W.):"
    table.rows[0].cells[1].text = "500.0 KGS"
    table.rows[1].cells[0].text = "Total Gross Weight (G.W.):"
    table.rows[1].cells[1].text = "550.0 KGS"
    table.rows[2].cells[0].text = "Total Measurement (CBM):"
    table.rows[2].cells[1].text = "2.5 CBM"
    table.rows[3].cells[0].text = "Packing Type & Marks:"
    table.rows[3].cells[1].text = "Type: 10 Wooden Pallets\n[Shipping Marks] GLOBAL TRADE INC."
    return doc

def create_proforma_invoice_docx(data):
    doc = Document()
    doc.add_heading('PROFORMA INVOICE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    create_common_table(doc, data)
    create_item_table(doc, data)
    doc.add_paragraph("\n[TERMS & CONDITIONS]")
    doc.add_paragraph("4. Remittance Instructions: Payment via T/T to Korea Bank (Acc: 987654321, SWIFT: KOBKRSSE).")
    return doc

# --- [7. 사이드바 구성] (✅새로운 네비게이션 구조 적용) ---
with st.sidebar:
    # 1) 국가별 수출입 데이터
    with st.expander("1) 해외진출 전략 허브", expanded=False):
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("시장동향", use_container_width=True, key="nav_cn_1"):
                st.switch_page("pages/macro_1.py")
        with col2:
            if st.button("전략분석", use_container_width=True, key="nav_cn_2"):
                st.switch_page("pages/micro_1.py")
        with col3:
            if st.button("규제진단", use_container_width=True, key="nav_cn_3"):
                st.switch_page("pages/mac_mic_1.py")

    # 2) SEO 서비스
    with st.expander("2) SEO 서비스", expanded=False):
        if st.button("바로가기", use_container_width=True, key="nav_news"):
            st.switch_page("pages/junghyun.py")

    # 3) AI 바이어 매칭 서비스
    with st.expander("3) AI 바이어 매칭 서비스", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            if st.button("바이어 찾기", use_container_width=True, key="nav_ai_1"):
                st.switch_page("pages/03_ai_chatbot.py")
        with col2:
            if st.button("전시회 일정", use_container_width=True, key="nav_ai_2"):
                st.switch_page("pages/buyer_maps.py")

    # 4) 환율 정보 확인
    with st.expander("4) 환율 정보 확인", expanded=False):
        if st.button("바로가기", use_container_width=True, key="nav_ex"):
            st.switch_page("pages/exchange_rate.py")

    # 5) 무역 서류 자동 완성
    with st.expander("5) 무역 서류 자동 완성", expanded=False):
        if st.button("바로가기", use_container_width=True, key="nav_fx"):
            st.switch_page("pages/auto_docs.py")

    # 로고 영역
    logo_path = "assets/logo.png"
    if os.path.exists(logo_path):
        logo_b64 = base64.b64encode(open(logo_path, "rb").read()).decode()
        st.markdown(
            f"""
            <div class="logo-box">
              <img class="logo-img" src="data:image/png;base64,{logo_b64}" alt="logo"/>
              <div class="small-muted" style="margin-top:8px; text-align:center;">
                KITA AX MASTER TEAM4
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.warning("⚠️ assets/logo.png 를 찾을 수 없습니다.")

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("---")

    # ✅ 홈으로 돌아가기 버튼 (맨 아래)
    if st.button("🏠 홈으로 돌아가기", key="go_home_btn", use_container_width=True):
        st.switch_page("dashboard.py")

# --- [8. 메인 화면 - Trade Master 2026] ---
st.title("🚢 환율 체크 및 차트 분석")

exchange_rates = st.session_state['exchange_rates']
use_realtime = st.session_state['use_realtime']
today_date = datetime.now().date()

# ✅ 제목과 버튼을 나란히 배치
header_col1, header_col2 = st.columns([4, 1])

with header_col1:
    st.subheader(f"주요 통화 환율 및 변동성 분석 ({'실시간' if use_realtime else '2026 시뮬레이션'})")

with header_col2:
    st.markdown("<br>", unsafe_allow_html=True)  # 버튼 위치 조정
    if st.button("🔄 최신 환율 반영", use_container_width=True, key="update_rates_btn"):
        with st.spinner("가져오는 중..."):
            updated_rates = get_realtime_exchange_rates()
            st.session_state['exchange_rates'] = updated_rates
            st.session_state['use_realtime'] = True
            st.success("✅ 최신 환율이 반영!")
            st.rerun()

# 환율 데이터 업데이트
exchange_rates = st.session_state['exchange_rates']
use_realtime = st.session_state['use_realtime']

currency_info = [
    ("🇺🇸 USD 미국", "USDKRW=X", exchange_rates['USD'], 1),
    ("🇯🇵 JPY 일본(100엔)", "JPYKRW=X", exchange_rates['JPY'], 100),
    ("🇪🇺 EUR 유럽", "EURKRW=X", exchange_rates['EUR'], 1),
    ("🇨🇳 CNY 중국", "CNYKRW=X", exchange_rates['CNY'], 1)
]

display_data = []
for name, ticker, base, mult in currency_info:
    hist_df = get_currency_history(ticker, base, mult, use_realtime, today_date)
    sparkline_values = hist_df['환율'].tail(15).tolist() if not hist_df.empty else [base] * 15
    np.random.seed(int(base))
    volatility = round(np.random.uniform(-1.8, 1.8), 2)
    display_data.append({
        "통화명": name,
        "매매기준율": base,
        "보낼 때 (1%)": base * 1.01,
        "받을 때 (1%)": base * 0.99,
        "시장 변동성 (%)": volatility,
        "최근 흐름 (Sparkline)": sparkline_values
    })

df_rates = pd.DataFrame(display_data)
st.dataframe(
    df_rates,
    column_config={
        "통화명": st.column_config.TextColumn("통화명", width="medium"),
        "매매기준율": st.column_config.NumberColumn("매매기준율", format="%.2f KRW"),
        "보낼 때 (1%)": st.column_config.NumberColumn("송금 보낼 때", format="%.2f"),
        "받을 때 (1%)": st.column_config.NumberColumn("송금 받을 때", format="%.2f"),
        "시장 변동성 (%)": st.column_config.ProgressColumn("시장 변동성 (%)", format="%.2f%%", min_value=-2.0, max_value=2.0),
        "최근 흐름 (Sparkline)": st.column_config.LineChartColumn("최근 흐름 (15일 추이)", width="medium")
    },
    hide_index=True,
    use_container_width=True
)

# --- [AI 인사이트 섹션] ---
st.markdown("### AI 실시간 시장 진단")
if st.button("AI에게 환율 데이터 분석 요청하기", use_container_width=True):
    with st.spinner("AI가 글로벌 시장 흐름을 분석 중입니다..."):
        # 위에서 정의한 함수 호출
        insight_text = get_ai_exchange_insight(df_rates)
        
        # 결과 출력 공간
        st.info(insight_text)
        st.caption("※ 본 분석은 데이터 기반 AI의 견해이며, 실제 거래 시에는 전문가와 상의하시기 바랍니다.")

st.divider()
st.subheader("실시간 환율 변동 시각화 (Financial Candlestick)")

chart_cols = st.columns(2)
for i, (label, ticker, base, mult) in enumerate(currency_info):
    with chart_cols[i % 2]:
        df_hist = get_currency_history(ticker, base, mult, use_realtime, today_date)
        if not df_hist.empty:
            st.plotly_chart(draw_candlestick_chart(df_hist, label), use_container_width=True)

# --- Footer ---
st.divider()
st.markdown("""
<div style='text-align: center; color: #718096; font-size: 0.9em;'>
    <p>Global E-commerce All In One Solution</p>
    <p>Developed by Seyeon Global Connect</p>
</div>
""", unsafe_allow_html=True)