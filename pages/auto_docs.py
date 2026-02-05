import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import base64
from typing import Dict, List, Tuple, Optional
import math
import requests
import os
from dotenv import load_dotenv

load_dotenv()

try:
    from openai import OpenAI
    api_key = os.getenv("OPENAI_API_KEY") or os.getenv("Open_api_key")
    client = OpenAI(api_key=api_key) if api_key else None
except:
    client = None

st.set_page_config(page_title="세연 글로벌 커넥트", page_icon="🚢", layout="wide", initial_sidebar_state="collapsed")

# ==================== 사이드바 네비게이션 ====================
st.markdown("""
    <style>
    [data-testid="stSidebarNav"] { display: none; }
    section[data-testid="stSidebar"] { 
        background: #ffffff !important;
        border-right: 1px solid #e5e7eb;
    }
    .stButton>button {
        background: #051161 !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: 8px;
        padding: 10px 14px;
        font-weight: 700;
        transition: 0.3s;
    }
    .stButton>button:hover {
        background: rgba(5,17,97,0.85) !important;
        box-shadow: 0 4px 12px rgba(5,17,97,0.3) !important;
    }
    .logo-box{
        background: rgba(255,255,255,0.6);
        border: 1px solid #e5e7eb;
        border-radius: 14px;
        padding: 14px 12px;
        margin-bottom: 10px;
        text-align:center;
    }
    .logo-img{
        max-width: 150px;
        width: 100%;
        height: auto;
        display:block;
        margin: 0 auto;
    }
    .small-muted{
        color:#64748b;
        font-size: 0.85rem;
        font-weight: 600;
        letter-spacing: 0.2px;
    }
    </style>
""", unsafe_allow_html=True)

with st.sidebar:
    # 1) 국가별 수출입 데이터
    with st.expander("1) 해외진출 전략 허브", expanded=False):
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("시장동향", use_container_width=True, key="nav_cn_1"):
                st.switch_page("pages/junghyun.py")
        with col2:
            if st.button("전략분석", use_container_width=True, key="nav_cn_2"):
                st.switch_page("pages/junghyun.py")
        with col3:
            if st.button("규제진단", use_container_width=True, key="nav_cn_3"):
                st.switch_page("pages/junghyun.py")

    # 2) SEO 서비스
    with st.expander("2) SEO 서비스", expanded=False):
        if st.button("바로가기", use_container_width=True, key="nav_news"):
            st.switch_page("pages/junghyun.py")

    # 3) AI 바이어 매칭 서비스
    with st.expander("3) AI 바이어 매칭 서비스", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            if st.button("바이어 찾기", use_container_width=True, key="nav_ai_1"):
                st.switch_page("pages/03_ai_chatbot.py")
        with col2:
            if st.button("전시회 일정", use_container_width=True, key="nav_ai_2"):
                st.switch_page("pages/buyer_maps.py")

    # 4) 환율 정보 확인
    with st.expander("4) 환율 정보 확인", expanded=False):
        if st.button("바로가기", use_container_width=True, key="nav_ex"):
            st.switch_page("pages/exchange_rate.py")

    # 5) 무역 서류 자동 완성
    with st.expander("5) 무역 서류 자동 완성", expanded=False):
        if st.button("바로가기", use_container_width=True, key="nav_fx"):
            st.switch_page("pages/auto_docs.py")

    # 로고 영역
    logo_path = "assets/logo.png"
    if os.path.exists(logo_path):
        logo_b64 = base64.b64encode(open(logo_path, "rb").read()).decode()
        st.markdown(
            f"""
            <div class="logo-box">
              <img class="logo-img" src="data:image/png;base64,{logo_b64}" alt="logo"/>
              <div class="small-muted" style="margin-top:8px; text-align:center;">
                KITA AX MASTER TEAM4
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            """
            <div class="logo-box">
              <div style="font-size:1.15rem; font-weight:900; color:#0f172a;">🚀 SEO Suite</div>
              <div class="small-muted" style="margin-top:6px;">KITA AX MASTER TEAM4</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("---")

    # 홈으로 돌아가기
    if st.button("🏠 홈으로 돌아가기", use_container_width=True, key="go_home_sidebar"):
        st.switch_page("dashboard.py")

# ==================== 사이드바 코드 끝 ====================

DESTINATION_PORTS = {
    "부산항": {"code": "KRPUS", "lat": 35.1040, "lon": 129.0403},
    "인천항": {"code": "KRINC", "lat": 37.4562, "lon": 126.7052},
    "광양항": {"code": "KRKWG", "lat": 34.9167, "lon": 127.6833}
}

ORIGIN_CITIES = ["구미", "청주", "화성", "수도권", "부산", "대구", "광주", "울산", "창원"]

KCCI_FREIGHT_RATES = {
    "부산항": {
        "LA (미서부)": {"20ft": 2800, "40ft": 3500, "40hc": 3700, "lcl_per_cbm": 145, "transit_days": 14},
        "롱비치 (미서부)": {"20ft": 2900, "40ft": 3600, "40hc": 3800, "lcl_per_cbm": 150, "transit_days": 15},
        "뉴욕 (미동부)": {"20ft": 4500, "40ft": 5800, "40hc": 6200, "lcl_per_cbm": 220, "transit_days": 28},
        "상하이": {"20ft": 450, "40ft": 650, "40hc": 720, "lcl_per_cbm": 55, "transit_days": 2},
        "함부르크 (유럽)": {"20ft": 4200, "40ft": 5500, "40hc": 6000, "lcl_per_cbm": 230, "transit_days": 35},
        "로테르담 (유럽)": {"20ft": 4300, "40ft": 5600, "40hc": 6100, "lcl_per_cbm": 235, "transit_days": 36},
        "싱가포르": {"20ft": 850, "40ft": 1200, "40hc": 1350, "lcl_per_cbm": 92, "transit_days": 7},
        "호치민": {"20ft": 680, "40ft": 950, "40hc": 1050, "lcl_per_cbm": 78, "transit_days": 5},
        "도쿄": {"20ft": 420, "40ft": 580, "40hc": 640, "lcl_per_cbm": 48, "transit_days": 3},
    },
    "인천항": {
        "상하이": {"20ft": 450, "40ft": 650, "40hc": 720, "lcl_per_cbm": 55, "transit_days": 2},
        "칭다오": {"20ft": 380, "40ft": 520, "40hc": 580, "lcl_per_cbm": 48, "transit_days": 2},
        "호치민": {"20ft": 700, "40ft": 980, "40hc": 1080, "lcl_per_cbm": 82, "transit_days": 6},
        "도쿄": {"20ft": 420, "40ft": 580, "40hc": 640, "lcl_per_cbm": 48, "transit_days": 3},
        "LA (미서부)": {"20ft": 2900, "40ft": 3600, "40hc": 3800, "lcl_per_cbm": 150, "transit_days": 16},
    },
    "광양항": {
        "상하이": {"20ft": 420, "40ft": 620, "40hc": 690, "lcl_per_cbm": 52, "transit_days": 3},
        "호치민": {"20ft": 720, "40ft": 1000, "40hc": 1100, "lcl_per_cbm": 85, "transit_days": 6},
        "싱가포르": {"20ft": 880, "40ft": 1250, "40hc": 1400, "lcl_per_cbm": 95, "transit_days": 8},
    }
}

INLAND_TRANSPORT_RATES = {
    "구미": {"부산항": {"5톤": 280000, "20ft": 420000, "40ft": 580000}, "인천항": {"5톤": 350000, "20ft": 520000, "40ft": 720000}, "광양항": {"5톤": 180000, "20ft": 280000, "40ft": 380000}},
    "청주": {"부산항": {"5톤": 320000, "20ft": 480000, "40ft": 660000}, "인천항": {"5톤": 180000, "20ft": 280000, "40ft": 380000}, "광양항": {"5톤": 280000, "20ft": 420000, "40ft": 580000}},
    "화성": {"부산항": {"5톤": 380000, "20ft": 560000, "40ft": 780000}, "인천항": {"5톤": 120000, "20ft": 180000, "40ft": 240000}, "광양항": {"5톤": 350000, "20ft": 520000, "40ft": 720000}},
    "수도권": {"부산항": {"5톤": 350000, "20ft": 520000, "40ft": 720000}, "인천항": {"5톤": 150000, "20ft": 220000, "40ft": 300000}, "광양항": {"5톤": 320000, "20ft": 480000, "40ft": 660000}},
    "부산": {"부산항": {"5톤": 80000, "20ft": 120000, "40ft": 160000}, "인천항": {"5톤": 420000, "20ft": 620000, "40ft": 860000}, "광양항": {"5톤": 180000, "20ft": 280000, "40ft": 380000}},
    "대구": {"부산항": {"5톤": 180000, "20ft": 280000, "40ft": 380000}, "인천항": {"5톤": 320000, "20ft": 480000, "40ft": 660000}, "광양항": {"5톤": 220000, "20ft": 340000, "40ft": 460000}},
    "광주": {"부산항": {"5톤": 280000, "20ft": 420000, "40ft": 580000}, "인천항": {"5톤": 380000, "20ft": 560000, "40ft": 780000}, "광양항": {"5톤": 120000, "20ft": 180000, "40ft": 240000}},
    "울산": {"부산항": {"5톤": 120000, "20ft": 180000, "40ft": 240000}, "인천항": {"5톤": 380000, "20ft": 560000, "40ft": 780000}, "광양항": {"5톤": 220000, "20ft": 340000, "40ft": 460000}},
    "창원": {"부산항": {"5톤": 100000, "20ft": 150000, "40ft": 200000}, "인천항": {"5톤": 400000, "20ft": 590000, "40ft": 820000}, "광양항": {"5톤": 180000, "20ft": 280000, "40ft": 380000}},
}

PORT_CHARGES = {
    "부산항": {
        "FCL": {
            "20ft": {"thc": 185000, "wharfage": 4429, "doc_fee": 50000, "handling_charge": 70000, "seal_fee": 8000, "container_tax": 1500},
            "40ft": {"thc": 280000, "wharfage": 8858, "doc_fee": 50000, "handling_charge": 100000, "seal_fee": 8000, "container_tax": 3000},
            "40hc": {"thc": 295000, "wharfage": 8858, "doc_fee": 50000, "handling_charge": 105000, "seal_fee": 8000, "container_tax": 3000},
        },
        "LCL": {"thc_per_cbm": 22000, "cfs_charge_per_cbm": 38000, "drayage_per_cbm": 15000, "minimum_charge": 200000}
    },
    "인천항": {
        "FCL": {
            "20ft": {"thc": 175000, "wharfage": 4200, "doc_fee": 45000, "handling_charge": 65000, "seal_fee": 8000, "container_tax": 1500},
            "40ft": {"thc": 265000, "wharfage": 8400, "doc_fee": 45000, "handling_charge": 95000, "seal_fee": 8000, "container_tax": 3000},
            "40hc": {"thc": 280000, "wharfage": 8400, "doc_fee": 45000, "handling_charge": 100000, "seal_fee": 8000, "container_tax": 3000},
        },
        "LCL": {"thc_per_cbm": 20000, "cfs_charge_per_cbm": 35000, "drayage_per_cbm": 14000, "minimum_charge": 180000}
    },
    "광양항": {
        "FCL": {
            "20ft": {"thc": 150000, "wharfage": 4000, "doc_fee": 40000, "handling_charge": 55000, "seal_fee": 7000, "container_tax": 1400},
            "40ft": {"thc": 230000, "wharfage": 8000, "doc_fee": 40000, "handling_charge": 80000, "seal_fee": 7000, "container_tax": 2800},
            "40hc": {"thc": 245000, "wharfage": 8000, "doc_fee": 40000, "handling_charge": 85000, "seal_fee": 7000, "container_tax": 2800},
        },
        "LCL": {"thc_per_cbm": 18000, "cfs_charge_per_cbm": 32000, "drayage_per_cbm": 12000, "minimum_charge": 160000}
    }
}

DEFAULT_EXCHANGE_RATE = 1450

INCOTERMS_LIST = ["EXW (Ex Works)", "FCA (Free Carrier)", "FAS (Free Alongside Ship)", "FOB (Free On Board)", "CFR (Cost and Freight)", "CIF (Cost, Insurance and Freight)", "CPT (Carriage Paid To)", "CIP (Carriage and Insurance Paid To)", "DAP (Delivered at Place)", "DPU (Delivered at Place Unloaded)", "DDP (Delivered Duty Paid)", "DAT (Delivered at Terminal - Legacy)"]

FTA_RATES = {"협정 미적용 (기본세율 8%)": 0.08, "한-미 FTA (KOR-USA)": 0.0, "한-EU FTA (KOR-EU)": 0.0, "한-중 FTA (KOR-CHINA)": 0.054, "한-아세안 FTA (KOR-ASEAN)": 0.0, "한-베트남 FTA": 0.0, "한-인도 CEPA": 0.025, "한-영 FTA (KOR-UK)": 0.0, "RCEP (역내포괄적경제동반자협정)": 0.03, "한-호주 FTA": 0.0, "한-캐나다 FTA": 0.0, "한-뉴질랜드 FTA": 0.0, "한-싱가포르 FTA": 0.0, "한-EFTA FTA": 0.0}

INSURANCE_LIST = ["선택 안함", "ICC(A) - 전위험담보 (0.8%)", "ICC(B) - 분손담보 (0.5%)", "ICC(C) - 분손부담보 (0.3%)", "ICC(Air) - 항공화물 (0.6%)", "TPL (Third Party Liability) (0.2%)"]
INSURANCE_RATES = {"선택 안함": 0, "ICC(A) - 전위험담보 (0.8%)": 0.008, "ICC(B) - 분손담보 (0.5%)": 0.005, "ICC(C) - 분손부담보 (0.3%)": 0.003, "ICC(Air) - 항공화물 (0.6%)": 0.006, "TPL (Third Party Liability) (0.2%)": 0.002}

PAYMENT_METHODS_PRESET = {"L/C at Sight (신용장 일람불)": {"fee": 0.012, "desc": "L/C at Sight"}, "L/C Usance 30일": {"fee": 0.015, "desc": "L/C Usance 30 days"}, "L/C Usance 60일": {"fee": 0.018, "desc": "L/C Usance 60 days"}, "T/T in Advance (사전송금)": {"fee": 0.001, "desc": "T/T in Advance"}, "T/T 30% Advance, 70% against B/L": {"fee": 0.003, "desc": "T/T 30% + 70% B/L"}, "D/P at Sight (지급인도조건)": {"fee": 0.005, "desc": "D/P at Sight"}, "D/A 30 Days (인수인도조건)": {"fee": 0.007, "desc": "D/A 30 Days"}, "직접 입력": {"fee": 0.005, "desc": "Custom"}}

st.markdown("""
<style>
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
    html, body, [class*="css"], .stApp { font-family: 'Pretendard', sans-serif; }
    .main { background: linear-gradient(135deg, #f5f7fa 0%, #e8ecf1 100%); }
    .stButton>button { width: 100%; border-radius: 8px; background: linear-gradient(90deg, #1a237e 0%, #283593 100%); color: white; font-weight: 600; height: 3.5em; border: none; font-size: 1rem; }
    .stButton>button:hover { transform: translateY(-2px); box-shadow: 0 4px 12px rgba(0,0,0,0.2); }
    .content-box { background: white; border-radius: 12px; padding: 25px; border: 1px solid #e5e7eb; margin-bottom: 20px; }
    .result-box { background: #f8f9fa; border-radius: 12px; padding: 20px; border: 1px solid #e5e7eb; }
    .savings-box { background: #e8f5e9; border-radius: 8px; padding: 12px; margin-top: 10px; }
</style>
""", unsafe_allow_html=True)

class ProfessionalDocGenerator:
    def __init__(self, data: Dict, seal_path=None, trade_mode="수출"):
        self.data = data
        self.seal_path = seal_path
        self.trade_mode = trade_mode
        self.today = datetime.now().strftime("%Y-%m-%d")
    
    def _set_cell_bg(self, cell, color="E8EAF6"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        tcPr.append(shading)
    
    def _format_header_cell(self, cell, label, value):
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run_lbl = p.add_run(f"{label}\n")
        run_lbl.font.size = Pt(8)
        run_lbl.font.bold = True
        run_lbl.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
        run_val = p.add_run(str(value))
        run_val.font.size = Pt(10)
    
    def _add_seal_to_cell(self, cell):
        if self.seal_path:
            try:
                if hasattr(self.seal_path, 'seek'):
                    self.seal_path.seek(0)
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(self.seal_path, width=Cm(2.5))
            except:
                pass
    
    def _add_seal_to_cell_simple(self, doc):
        if self.seal_path:
            try:
                if hasattr(self.seal_path, 'seek'):
                    self.seal_path.seek(0)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run()
                run.add_picture(self.seal_path, width=Cm(2.5))
            except:
                pass
    
    def create_commercial_invoice(self):
        doc = Document()
        title = doc.add_heading('COMMERCIAL INVOICE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        headers = [("Exporter (Seller):", f"{self.data['shipper']}\n{self.data.get('shipper_addr','')}", "Invoice No. & Date:", f"{self.data['inv_no']} / {self.today}"),
                   ("Consignee (Buyer):", f"{self.data['consignee']}\n{self.data.get('consignee_addr','')}", "L/C Issuing Bank:", self.data.get('lc_bank', 'N/A')),
                   ("Notify Party:", self.data.get('notify_party', 'Same as Consignee'), "Port of Loading / Discharge:", f"{self.data['pol']} / {self.data['pod']}"),
                   ("Vessel / Flight No:", self.data.get('vessel', 'TBA'), "Terms of Delivery & Payment:", f"{self.data['incoterms']} / {self.data['payment']}")]
        for r, (l1, v1, l2, v2) in enumerate(headers):
            self._format_header_cell(table.rows[r].cells[0], l1, v1)
            self._format_header_cell(table.rows[r].cells[1], l2, v2)
        doc.add_paragraph()
        item_header = doc.add_paragraph("[ITEM DESCRIPTION]")
        item_header.runs[0].font.bold = True
        item_table = doc.add_table(rows=1, cols=5)
        item_table.style = 'Table Grid'
        cols = ["No.", "Description of Goods", "Quantity", "Unit Price", "Amount"]
        for i, col in enumerate(cols):
            cell = item_table.rows[0].cells[i]
            self._set_cell_bg(cell, "E8EAF6")
            cell.text = col
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        qty = self.data.get('qty', 100)
        unit_price = self.data.get('unit_price', 50.00)
        total = qty * unit_price
        currency = self.data.get('currency', 'USD')
        row = item_table.add_row().cells
        row[0].text = "1"
        row[1].text = f"{self.data['item_desc']}\n(HS Code: {self.data['hs_code']})"
        row[2].text = f"{qty} {self.data.get('unit', 'PCS')}"
        row[3].text = f"{currency} {unit_price:,.2f}"
        row[4].text = f"{currency} {total:,.2f}"
        total_row = item_table.add_row().cells
        total_row[3].text = "TOTAL:"
        total_row[3].paragraphs[0].runs[0].font.bold = True
        total_row[4].text = f"{currency} {total:,.2f}"
        total_row[4].paragraphs[0].runs[0].font.bold = True
        doc.add_paragraph()
        clause_header = doc.add_paragraph("[SPECIAL CLAUSES]")
        clause_header.runs[0].font.bold = True
        ai_clauses = self.data.get('ai_clauses', 'Standard Trade Terms Apply per INCOTERMS 2020.')
        doc.add_paragraph(ai_clauses)
        doc.add_paragraph("\n")
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.style = 'Table Grid'
        c1 = sig_table.rows[0].cells[0]
        c1.paragraphs[0].add_run("Authorized Signature").font.bold = True
        c1.add_paragraph(f"\n{self.data['shipper'].split(chr(10))[0]}")
        c1.add_paragraph("\n")
        self._add_seal_to_cell(c1)
        c1.add_paragraph("_________________________")
        c1.add_paragraph("(Signature & Company Stamp)")
        c1.add_paragraph(f"Date: {self.today}")
        return doc

    def create_packing_list(self):
        doc = Document()
        title = doc.add_heading('PACKING LIST', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        headers = [("Shipper:", f"{self.data['shipper']}\n{self.data.get('shipper_addr','')}", "P/L No. & Date:", f"{self.data['inv_no']}-PL / {self.today}"),
                   ("Consignee:", f"{self.data['consignee']}\n{self.data.get('consignee_addr','')}", "POL / POD:", f"{self.data['pol']} / {self.data['pod']}"),
                   ("Notify Party:", self.data.get('notify_party', 'Same as Consignee'), "Vessel / Flight:", self.data.get('vessel', 'TBA'))]
        for r, (l1, v1, l2, v2) in enumerate(headers):
            self._format_header_cell(table.rows[r].cells[0], l1, v1)
            self._format_header_cell(table.rows[r].cells[1], l2, v2)
        doc.add_paragraph()
        pack_header = doc.add_paragraph("[PACKING DETAILS]")
        pack_header.runs[0].font.bold = True
        pack_table = doc.add_table(rows=1, cols=6)
        pack_table.style = 'Table Grid'
        cols = ["No.", "Description", "Quantity", "N.W (KG)", "G.W (KG)", "Measurement (CBM)"]
        for i, col in enumerate(cols):
            cell = pack_table.rows[0].cells[i]
            self._set_cell_bg(cell, "E8EAF6")
            cell.text = col
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row = pack_table.add_row().cells
        row[0].text = "1"
        row[1].text = f"{self.data['item_desc']}\n(HS Code: {self.data['hs_code']})"
        row[2].text = f"{self.data['qty']} {self.data.get('unit', 'PCS')}"
        row[3].text = f"{self.data.get('nw', '-')}"
        row[4].text = f"{self.data.get('gw', '-')}"
        row[5].text = f"{self.data.get('cbm', '-')}"
        total_row = pack_table.add_row().cells
        total_row[1].text = "TOTAL"
        total_row[1].paragraphs[0].runs[0].font.bold = True
        total_row[2].text = f"{self.data['qty']} {self.data.get('unit', 'PCS')}"
        total_row[3].text = f"{self.data.get('nw', '-')}"
        total_row[4].text = f"{self.data.get('gw', '-')}"
        total_row[5].text = f"{self.data.get('cbm', '-')}"
        doc.add_paragraph()
        pack_info = doc.add_paragraph("[PACKING INFORMATION]")
        pack_info.runs[0].font.bold = True
        doc.add_paragraph(f"Packing Type: {self.data.get('packing_type', 'Standard Export Packing')}")
        doc.add_paragraph(f"Shipping Marks: {self.data.get('marks', 'N/M')}")
        doc.add_paragraph("\n")
        sig_p = doc.add_paragraph()
        sig_p.add_run("Prepared by: ").font.bold = True
        sig_p.add_run(self.data['shipper'].split('\n')[0])
        self._add_seal_to_cell_simple(doc)
        return doc
    
    def create_sales_contract(self):
        doc = Document()
        title = doc.add_heading('SALES CONTRACT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Contract No.: {self.data['inv_no']}-SC")
        doc.add_paragraph(f"Date: {self.today}")
        doc.add_paragraph()
        parties = doc.add_paragraph("[PARTIES]")
        parties.runs[0].font.bold = True
        party_table = doc.add_table(rows=2, cols=2)
        party_table.style = 'Table Grid'
        self._format_header_cell(party_table.rows[0].cells[0], "SELLER:", self.data['shipper'])
        self._format_header_cell(party_table.rows[0].cells[1], "BUYER:", self.data['consignee'])
        self._format_header_cell(party_table.rows[1].cells[0], "Address:", self.data.get('shipper_addr', ''))
        self._format_header_cell(party_table.rows[1].cells[1], "Address:", self.data.get('consignee_addr', ''))
        doc.add_paragraph()
        terms_header = doc.add_paragraph("[TERMS AND CONDITIONS]")
        terms_header.runs[0].font.bold = True
        terms_table = doc.add_table(rows=6, cols=2)
        terms_table.style = 'Table Grid'
        qty = self.data.get('qty', 100)
        unit_price = self.data.get('unit_price', 50.00)
        total = qty * unit_price
        currency = self.data.get('currency', 'USD')
        terms_data = [("1. Commodity:", f"{self.data['item_desc']} (HS Code: {self.data['hs_code']})"), ("2. Quantity:", f"{qty} {self.data.get('unit', 'PCS')}"), ("3. Unit Price:", f"{currency} {unit_price:,.2f}"), ("4. Total Amount:", f"{currency} {total:,.2f}"), ("5. Terms of Delivery:", self.data['incoterms']), ("6. Payment Terms:", self.data['payment'])]
        for r, (label, value) in enumerate(terms_data):
            terms_table.rows[r].cells[0].text = label
            terms_table.rows[r].cells[0].paragraphs[0].runs[0].font.bold = True
            terms_table.rows[r].cells[1].text = value
        doc.add_paragraph()
        ship_header = doc.add_paragraph("[SHIPMENT DETAILS]")
        ship_header.runs[0].font.bold = True
        doc.add_paragraph(f"Port of Loading: {self.data['pol']}")
        doc.add_paragraph(f"Port of Discharge: {self.data['pod']}")
        doc.add_paragraph(f"Vessel/Flight: {self.data.get('vessel', 'TBA')}")
        doc.add_paragraph()
        clause_header = doc.add_paragraph("[SPECIAL CLAUSES]")
        clause_header.runs[0].font.bold = True
        ai_clauses = self.data.get('ai_clauses', 'Standard Trade Terms Apply per INCOTERMS 2020.')
        doc.add_paragraph(ai_clauses)
        doc.add_paragraph("\n")
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.style = 'Table Grid'
        c1 = sig_table.rows[0].cells[0]
        c1.paragraphs[0].add_run("SELLER").font.bold = True
        c1.add_paragraph(f"\n{self.data['shipper'].split(chr(10))[0]}")
        c1.add_paragraph("\n")
        self._add_seal_to_cell(c1)
        c1.add_paragraph("_________________________")
        c1.add_paragraph("(Authorized Signature & Stamp)")
        c1.add_paragraph(f"Date: {self.today}")
        c2 = sig_table.rows[0].cells[1]
        c2.paragraphs[0].add_run("BUYER").font.bold = True
        c2.add_paragraph(f"\n{self.data['consignee'].split(chr(10))[0]}")
        c2.add_paragraph("\n\n\n_________________________")
        c2.add_paragraph("(Authorized Signature)")
        c2.add_paragraph(f"Date: ________________")
        return doc
    
    def create_proforma_invoice(self):
        doc = Document()
        title = doc.add_heading('PROFORMA INVOICE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notice = doc.add_paragraph()
        notice.add_run("* This is a Proforma Invoice for reference purposes only and is not a demand for payment.").font.italic = True
        doc.add_paragraph()
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        headers = [("Exporter:", f"{self.data['shipper']}", "P/I No. & Date:", f"{self.data['inv_no']}-PI / {self.today}"), ("Importer:", f"{self.data['consignee']}", "Validity:", "30 days from date of issue"), ("POL / POD:", f"{self.data['pol']} / {self.data['pod']}", "Terms:", f"{self.data['incoterms']} / {self.data['payment']}"), ("Vessel:", self.data.get('vessel', 'TBA'), "Payment:", self.data['payment'])]
        for r, (l1, v1, l2, v2) in enumerate(headers):
            self._format_header_cell(table.rows[r].cells[0], l1, v1)
            self._format_header_cell(table.rows[r].cells[1], l2, v2)
        doc.add_paragraph()
        item_header = doc.add_paragraph("[ITEM DESCRIPTION]")
        item_header.runs[0].font.bold = True
        item_table = doc.add_table(rows=3, cols=5)
        item_table.style = 'Table Grid'
        cols = ["No.", "Description", "Quantity", "Unit Price", "Amount"]
        for i, col in enumerate(cols):
            cell = item_table.rows[0].cells[i]
            self._set_cell_bg(cell, "D9E2F3")
            cell.text = col
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        qty = self.data.get('qty', 100)
        unit_price = self.data.get('unit_price', 50.00)
        total = qty * unit_price
        currency = self.data.get('currency', 'USD')
        item_table.rows[1].cells[0].text = "1"
        item_table.rows[1].cells[1].text = f"{self.data.get('item_desc', 'Goods')}\n(HS: {self.data.get('hs_code', '0000.00')})"
        item_table.rows[1].cells[2].text = f"{qty} {self.data.get('unit', 'PCS')}"
        item_table.rows[1].cells[3].text = f"{currency} {unit_price:,.2f}"
        item_table.rows[1].cells[4].text = f"{currency} {total:,.2f}"
        for i in range(4):
            item_table.rows[2].cells[i].text = "ESTIMATED TOTAL"
            self._set_cell_bg(item_table.rows[2].cells[i], "F2F2F2")
        item_table.rows[2].cells[4].text = f"{currency} {total:,.2f}"
        doc.add_paragraph()
        doc.add_paragraph("* This quotation is valid for 30 days from the date of issue.")
        doc.add_paragraph("* Final Commercial Invoice will be issued upon shipment.")
        return doc
    
    def create_purchase_order(self):
        doc = Document()
        title = doc.add_heading('PURCHASE ORDER', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"P/O No.: {self.data['inv_no']}-PO")
        doc.add_paragraph(f"Date: {self.today}")
        doc.add_paragraph()
        buyer_header = doc.add_paragraph("[BUYER INFORMATION]")
        buyer_header.runs[0].font.bold = True
        doc.add_paragraph(f"Company: {self.data['consignee'].split(chr(10))[0]}")
        doc.add_paragraph(f"Address: {self.data.get('consignee_addr', '')}")
        doc.add_paragraph()
        seller_header = doc.add_paragraph("[SELLER INFORMATION]")
        seller_header.runs[0].font.bold = True
        doc.add_paragraph(f"Company: {self.data['shipper'].split(chr(10))[0]}")
        doc.add_paragraph(f"Address: {self.data.get('shipper_addr', '')}")
        doc.add_paragraph()
        order_header = doc.add_paragraph("[ORDER DETAILS]")
        order_header.runs[0].font.bold = True
        order_table = doc.add_table(rows=1, cols=5)
        order_table.style = 'Table Grid'
        cols = ["No.", "Item Description", "Quantity", "Unit Price", "Amount"]
        for i, col in enumerate(cols):
            cell = order_table.rows[0].cells[i]
            self._set_cell_bg(cell, "FFF3E0")
            cell.text = col
            cell.paragraphs[0].runs[0].font.bold = True
        qty = self.data.get('qty', 100)
        unit_price = self.data.get('unit_price', 50.00)
        total = qty * unit_price
        currency = self.data.get('currency', 'USD')
        row = order_table.add_row().cells
        row[0].text = "1"
        row[1].text = f"{self.data['item_desc']}\n(HS Code: {self.data['hs_code']})"
        row[2].text = f"{qty} {self.data.get('unit', 'PCS')}"
        row[3].text = f"{currency} {unit_price:,.2f}"
        row[4].text = f"{currency} {total:,.2f}"
        total_row = order_table.add_row().cells
        total_row[3].text = "TOTAL:"
        total_row[3].paragraphs[0].runs[0].font.bold = True
        total_row[4].text = f"{currency} {total:,.2f}"
        doc.add_paragraph()
        terms_header = doc.add_paragraph("[TERMS]")
        terms_header.runs[0].font.bold = True
        doc.add_paragraph(f"Delivery Terms: {self.data['incoterms']}")
        doc.add_paragraph(f"Payment Terms: {self.data['payment']}")
        doc.add_paragraph(f"Shipment From: {self.data['pol']} To: {self.data['pod']}")
        doc.add_paragraph()
        sig_p = doc.add_paragraph()
        sig_p.add_run("Authorized by: ").font.bold = True
        sig_p.add_run("_________________________")
        self._add_seal_to_cell_simple(doc)
        return doc
    
    def create_lc_application(self):
        doc = Document()
        title = doc.add_heading('APPLICATION FOR DOCUMENTARY CREDIT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Date: {self.today}")
        doc.add_paragraph(f"To: {self.data.get('lc_bank', '[Issuing Bank Name]')}")
        doc.add_paragraph()
        app_header = doc.add_paragraph("[CREDIT DETAILS]")
        app_header.runs[0].font.bold = True
        app_table = doc.add_table(rows=8, cols=2)
        app_table.style = 'Table Grid'
        qty = self.data.get('qty', 100)
        unit_price = self.data.get('unit_price', 50.00)
        total = qty * unit_price
        currency = self.data.get('currency', 'USD')
        app_data = [("Form of Credit:", "Irrevocable"), ("Beneficiary:", f"{self.data['shipper']}"), ("Applicant:", f"{self.data['consignee']}"), ("Amount:", f"{currency} {total:,.2f}"), ("Expiry Date:", "60 days from issue"), ("Port of Loading:", self.data['pol']), ("Port of Discharge:", self.data['pod']), ("Terms:", self.data['incoterms'])]
        for r, (label, value) in enumerate(app_data):
            app_table.rows[r].cells[0].text = label
            app_table.rows[r].cells[0].paragraphs[0].runs[0].font.bold = True
            self._set_cell_bg(app_table.rows[r].cells[0], "E3F2FD")
            app_table.rows[r].cells[1].text = value
        doc.add_paragraph()
        goods_header = doc.add_paragraph("[DESCRIPTION OF GOODS]")
        goods_header.runs[0].font.bold = True
        doc.add_paragraph(f"{self.data['item_desc']}")
        doc.add_paragraph(f"HS Code: {self.data['hs_code']}")
        doc.add_paragraph(f"Quantity: {qty} {self.data.get('unit', 'PCS')}")
        doc.add_paragraph()
        docs_header = doc.add_paragraph("[DOCUMENTS REQUIRED]")
        docs_header.runs[0].font.bold = True
        doc.add_paragraph("1. Commercial Invoice (Original x 3)")
        doc.add_paragraph("2. Packing List (Original x 3)")
        doc.add_paragraph("3. Full set of Clean On Board Bill of Lading")
        doc.add_paragraph("4. Certificate of Origin")
        doc.add_paragraph("5. Insurance Certificate (if CIF/CIP)")
        doc.add_paragraph()
        sig_p = doc.add_paragraph()
        sig_p.add_run("Applicant's Signature: ").font.bold = True
        sig_p.add_run("_________________________")
        self._add_seal_to_cell_simple(doc)
        return doc

def get_exchange_rate():
    try:
        import yfinance as yf
        data = yf.download("USDKRW=X", period="1d", progress=False)
        if not data.empty:
            return float(data['Close'].iloc[-1])
    except:
        pass
    return DEFAULT_EXCHANGE_RATE

def calculate_logistics_cost(origin, dest_port, destination, shipment_mode, container_type, cbm, include_inland=True):
    exchange_rate = get_exchange_rate()
    freight_data = KCCI_FREIGHT_RATES.get(dest_port, {}).get(destination, {})
    if not freight_data:
        freight_data = {"20ft": 2800, "40ft": 3500, "40hc": 3700, "lcl_per_cbm": 145, "transit_days": 14}
    if shipment_mode == "FCL":
        ocean_freight_usd = freight_data.get(container_type.lower(), freight_data.get("20ft", 2800))
    else:
        ocean_freight_usd = cbm * freight_data.get("lcl_per_cbm", 145)
    ocean_freight_krw = ocean_freight_usd * exchange_rate
    port_data = PORT_CHARGES.get(dest_port, PORT_CHARGES["부산항"])
    if shipment_mode == "FCL":
        container_charges = port_data["FCL"].get(container_type.lower(), port_data["FCL"]["20ft"])
        port_charges_detail = container_charges.copy()
        port_charges_krw = sum(container_charges.values())
    else:
        lcl_data = port_data["LCL"]
        port_charges_detail = {"thc": int(lcl_data["thc_per_cbm"] * cbm), "cfs_charge": int(lcl_data["cfs_charge_per_cbm"] * cbm), "drayage": int(lcl_data["drayage_per_cbm"] * cbm)}
        port_charges_krw = max(sum(port_charges_detail.values()), lcl_data["minimum_charge"])
    inland_cost_krw = 0
    if include_inland and origin and origin in INLAND_TRANSPORT_RATES:
        inland_data = INLAND_TRANSPORT_RATES[origin].get(dest_port, {})
        if shipment_mode == "FCL":
            truck_key = "40ft" if "40" in container_type else "20ft"
            inland_cost_krw = inland_data.get(truck_key, 420000)
        else:
            inland_cost_krw = inland_data.get("5톤", 280000)
    total_krw = ocean_freight_krw + port_charges_krw + inland_cost_krw
    return {"ocean_freight_usd": ocean_freight_usd, "ocean_freight_krw": ocean_freight_krw, "port_charges_krw": port_charges_krw, "port_charges_detail": port_charges_detail, "inland_cost_krw": inland_cost_krw, "total_krw": total_krw, "total_usd": total_krw / exchange_rate, "exchange_rate": exchange_rate, "transit_days": freight_data.get("transit_days", 14)}

def get_document_bytes(doc):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def fill_custom_template(template_file, data, seal_path=None):
    template_file.seek(0) 
    doc = Document(template_file)
    
    replacements = {
        "{{shipper}}": str(data.get('shipper', '')),
        "{{shipper_addr}}": str(data.get('shipper_addr', '')),
        "{{consignee}}": str(data.get('consignee', '')),
        "{{consignee_addr}}": str(data.get('consignee_addr', '')),
        "{{notify}}": str(data.get('notify_party', data.get('notify', ''))),
        "{{notify_party}}": str(data.get('notify_party', data.get('notify', ''))),
        "{{inv_no}}": str(data.get('inv_no', '')),
        "{{date}}": datetime.now().strftime("%Y-%m-%d"),
        "{{lc_bank}}": str(data.get('lc_bank', '')),
        "{{pol}}": str(data.get('pol', '')),
        "{{pod}}": str(data.get('pod', '')),
        "{{vessel}}": str(data.get('vessel', '')),
        "{{terms}}": str(data.get('incoterms', '')),
        "{{incoterms}}": str(data.get('incoterms', '')),
        "{{payment}}": str(data.get('payment', '')),
        "{{item_name}}": str(data.get('item_desc', '')),
        "{{item_desc}}": str(data.get('item_desc', '')),
        "{{hs_code}}": str(data.get('hs_code', '')),
        "{{qty}}": str(data.get('qty', '')),
        "{{unit}}": str(data.get('unit', '')),
        "{{currency}}": str(data.get('currency', '')),
        "{{price}}": f"{data.get('unit_price', 0):,.2f}",
        "{{unit_price}}": f"{data.get('unit_price', 0):,.2f}",
        "{{amount}}": f"{data.get('total_amount', 0):,.2f}",
        "{{total_amount}}": f"{data.get('total_amount', 0):,.2f}",
        "{{nw}}": str(data.get('nw', '')),
        "{{gw}}": str(data.get('gw', '')),
        "{{cbm}}": str(data.get('cbm', '')),
        "{{packing_type}}": str(data.get('packing_type', '')),
        "{{marks}}": str(data.get('marks', '')),
        "{{clauses}}": str(data.get('ai_clauses', '')),
        "{{fta}}": str(data.get('fta', '')),
        "{{insurance}}": str(data.get('insurance', ''))
    }

    def replace_text_in_paragraph(paragraph):
        for key, value in replacements.items():
            if key in paragraph.text:
                replaced = False
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        replaced = True
                
                if not replaced:
                    paragraph.text = paragraph.text.replace(key, value)

    for para in doc.paragraphs:
        replace_text_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para)

    if seal_path:
        inserted = False
        if hasattr(seal_path, 'seek'):
            seal_path.seek(0)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "{{seal}}" in para.text:
                            para.text = para.text.replace("{{seal}}", "") 
                            run = para.add_run()
                            if hasattr(seal_path, 'seek'):
                                seal_path.seek(0)
                            run.add_picture(seal_path, width=Cm(2.5))
                            inserted = True
                            
        if not inserted:
            try:
                if hasattr(seal_path, 'seek'):
                    seal_path.seek(0)
                doc.add_paragraph("\n")
                if doc.tables:
                    last_table = doc.tables[-1]
                    cell = last_table.rows[0].cells[0]
                    p = cell.add_paragraph()
                else:
                    p = doc.add_paragraph()
                
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(seal_path, width=Cm(2.5))
            except:
                pass
                
    return doc

if 'trade_mode' not in st.session_state:
    st.session_state.trade_mode = "수출"
if 'logistics_result' not in st.session_state:
    st.session_state.logistics_result = None
if 'generated_docs' not in st.session_state:
    st.session_state['generated_docs'] = {}
if 'ai_analysis_result' not in st.session_state:
    st.session_state['ai_analysis_result'] = None
if 'doc_data' not in st.session_state:
    st.session_state['doc_data'] = None
if 'selected_docs_list' not in st.session_state:
    st.session_state['selected_docs_list'] = []
if 'show_results' not in st.session_state:
    st.session_state['show_results'] = False

st.title("🚢 무역 서류 자동 완성 서비스")

picked = st.pills("", ["수출", "수입", "물류비계산"], default=st.session_state.trade_mode)
st.session_state.trade_mode = picked
trade_mode = st.session_state.trade_mode

if trade_mode == "물류비계산":
    with st.container(border=True):
        st.markdown("#### 물류비 견적 조건 설정")
        
        shipment_mode_full = st.radio("운송 방식", ["FCL (Full Container Load)", "LCL (Less than Container Load)"], horizontal=True)
        shipment_mode = "FCL" if "FCL" in shipment_mode_full else "LCL"
        
        st.divider()
        
        col1, col2 = st.columns(2)
        with col1:
            departure_port = st.selectbox("출발항", list(DESTINATION_PORTS.keys()))
            available_destinations = list(KCCI_FREIGHT_RATES.get(departure_port, {}).keys())
            if not available_destinations:
                available_destinations = ["LA (미서부)", "상하이", "싱가포르"]
            arrival_port = st.selectbox("도착항 (최종 목적지)", available_destinations)
        
        with col2:
            if shipment_mode == "FCL":
                container_type = st.selectbox("컨테이너 타입", ["20ft", "40ft", "40hc"])
                cbm = st.number_input("용적 (CBM)", min_value=0.1, max_value=100.0, value=10.0, step=0.5)
            else:
                container_type = "LCL"
                st.info("LCL은 용적(CBM)으로 계산됩니다")
                cbm = st.number_input("용적 (CBM)", min_value=0.1, max_value=100.0, value=1.0, step=0.1)
        
        st.divider()
        
        include_inland = st.checkbox("내륙 운송비 포함", value=False)
        origin_city = None
        if include_inland:
            origin_city = st.selectbox("출발지 (공장/창고)", ORIGIN_CITIES)
            
        if st.button("견적 산출", use_container_width=True):
            result = calculate_logistics_cost(origin=origin_city if include_inland else None, dest_port=departure_port, destination=arrival_port, shipment_mode=shipment_mode, container_type=container_type, cbm=cbm, include_inland=include_inland)
            
            result['shipment_mode_saved'] = shipment_mode
            result['container_type_saved'] = container_type
            result['departure_port_saved'] = departure_port
            result['arrival_port_saved'] = arrival_port
            
            st.session_state.logistics_result = result

    if st.session_state.logistics_result:
        result = st.session_state.logistics_result
        ex_rate = result['exchange_rate']
        
        mode_disp = result.get('shipment_mode_saved', shipment_mode)
        cont_disp = result.get('container_type_saved', container_type)
        dep_disp = result.get('departure_port_saved', departure_port)
        arr_disp = result.get('arrival_port_saved', arrival_port)
        
        st.markdown("### 견적 결과")
        
        with st.container(border=True):
            st.markdown(f"""
            <div style="text-align: center; padding-bottom: 20px; border-bottom: 1px dashed #ddd; margin-bottom: 20px;">
                <span style="font-size: 1rem; color: #666;">총 물류비</span><br>
                <span style="font-size: 2rem; font-weight: 800; color: #1a237e;">₩{result['total_krw']:,.0f}</span>
                <span style="font-size: 1.2rem; color: #555; margin-left: 8px;">(${result['total_usd']:,.2f})</span>
                <div style="font-size: 0.85rem; color: #888; margin-top: 5px;">(VAT 별도 / 환율: ₩{ex_rate:,.0f})</div>
            </div>
            """, unsafe_allow_html=True)

            col_left, col_right = st.columns(2)
            
            with col_left:
                st.markdown("**해상 운임 상세**")
                st.write(f"선택 항만: {dep_disp}")
                st.write(f"항로: {dep_disp}-{arr_disp.split()[0] if arr_disp else 'N/A'}")
                st.write(f"운송방식: {mode_disp}")
                st.write(f"컨테이너: {cont_disp}")
                st.write(f"해상운임: ${result['ocean_freight_usd']:,.2f} (₩{result['ocean_freight_krw']:,.0f})")
                
                inland_label = f"₩{result['inland_cost_krw']:,}" if include_inland else "₩0"
                st.write(f"내륙운송: {inland_label}")

            with col_right:
                st.markdown("**부대비용 상세**")
                st.write(f"부대비용 합계: ₩{result['port_charges_krw']:,}")
                
                charges = result.get('port_charges_detail', {})
                charges_labels = {"thc": "THC", "wharfage": "Wharfage", "doc_fee": "Doc Fee", "handling_charge": "H/C", "seal_fee": "Seal Fee", "container_tax": "Container Tax", "cfs_charge": "CFS Charge", "drayage": "Drayage"}
                
                for key, value in charges.items():
                    label = charges_labels.get(key, key)
                    st.caption(f"• {label}: ₩{value:,}")
                
                market_avg = result['port_charges_krw'] * 1.15
                savings = market_avg - result['port_charges_krw']
                savings_pct = (savings / market_avg * 100) if market_avg > 0 else 0
                
                st.markdown(f"""
                <div style="background-color: #f1f8e9; padding: 12px; border-radius: 8px; text-align: center; color: #2e7d32; font-size: 0.9rem; margin-top: 20px; margin-bottom: 5px; width: 100%;">
                    시장 평균 대비 ₩{savings:,.0f} 절감 ({savings_pct:.1f}%)
                </div>
                """, unsafe_allow_html=True)
else:
    with st.form("trade_form"):
        st.markdown("#### 서류 유형 선택")
        doc_options = ["Commercial Invoice", "Packing List", "Sales Contract", "Proforma Invoice", "L/C Application", "Purchase Order"]
        default_sel = ["Commercial Invoice", "Packing List", "Sales Contract"] if trade_mode == "수출" else ["Purchase Order", "Proforma Invoice"]
        selected_docs = st.multiselect("서류 유형을 선택해주세요", options=doc_options, default=default_sel, label_visibility="collapsed")
        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
        st.markdown("#### 수출입 당사자 정보")
        c1, c2, c3 = st.columns(3)
        if trade_mode == "수출":
            with c1:
                shipper = st.text_area("**Shipper/Seller**", "SeyeonGlobal Ltd.,")
                shipper_tel = st.text_input("**Tel (Shipper)**", "02-1234-5678")
                pol = st.text_input("**선적항(POL)**", "Busan, Korea")
            with c2:
                consignee = st.text_area("**Consignee**", "SeyeonGlobal Ltd.\n123 Samsung, Gangnam-gu,\nSeoul, Korea")
                consignee_tel = st.text_input("**Tel (Consignee)**", "02-1234-5678")
                pod = st.text_input("**도착항(POD)**", "LA, USA")
            with c3:
                notify_party = st.text_area("**Notify Party**", "Same as Consignee")
                inv_no = st.text_input("**서류 번호**", f"INV-{datetime.now().strftime('%Y%m%d')}-001")
                vessel = st.text_input("**선박/항공편명**", "TBA")
        else:
            with c1:
                shipper = st.text_area("**Supplier/Shipper**", "SeyeonGlobal Ltd.,")
                shipper_tel = st.text_input("**Tel (Supplier)**", "02-1234-5678")
                pol = st.text_input("**선적항(POL)**", "LA, USA")
            with c2:
                consignee = st.text_area("**Importer/Buyer**", "SeyeonGlobal Ltd.\n123 Samsung, Gangnam-gu,\nSeoul, Korea")
                consignee_tel = st.text_input("**Tel (Importer)**", "02-1234-5678")
                pod = st.text_input("**도착항(POD)**", "Busan, Korea")
            with c3:
                notify_party = st.text_area("**Notify Party**", "Same as Consignee")
                inv_no = st.text_input("**서류 번호**", f"INV-{datetime.now().strftime('%Y%m%d')}-001")
                vessel = st.text_input("**선박/항공편명**", "TBA")
        st.divider()
        st.markdown("#### 품목 정보")
        cc1, cc2, cc3 = st.columns(3)
        with cc1:
            item_desc = st.text_input("**품목명 (Description)**", "Electronic Components")
            hs_code = st.text_input("**HS Code**", "8517.62")
        with cc2:
            qty = st.number_input("**수량 (Quantity)**", value=100, min_value=1)
            unit = st.selectbox("**단위**", ["PCS", "SET", "KG", "MT", "CTN", "ROLL", "EA"])
        with cc3:
            selected_currency = st.selectbox("**통화**", ["USD", "EUR", "JPY", "CNY"])
            unit_price = st.number_input(f"**단가** ({selected_currency})", value=50.0, min_value=0.0, format="%.2f")
        st.divider()
        st.markdown("#### 인코텀즈 거래 조건")
        pc1, pc2, pc3 = st.columns(3)
        with pc1:
            transport_mode = st.radio("**운송수단**", ["해상(SEA)", "항공(AIR)"], horizontal=True)
            lc_bank = st.text_input("**L/C 개설은행**", "N/A")
        with pc2:
            incoterms = st.selectbox("**Incoterms**", INCOTERMS_LIST, index=3)
            selected_fta = st.selectbox("**FTA 협정세율**", list(FTA_RATES.keys()))
        with pc3:
            insurance_type = st.selectbox("**적하보험**", INSURANCE_LIST)
            
            payment_options = ["직접 입력 (아래 텍스트창 사용)"] + list(PAYMENT_METHODS_PRESET.keys())
            selected_payment = st.selectbox("**결제 방식**", payment_options)
            
            payment_custom_input = st.text_input("상세 조건 (직접 입력)", placeholder="내용을 입력하면 목록 선택보다 우선 적용됩니다.")
            
            if payment_custom_input.strip():
                payment_method = payment_custom_input
            elif "직접 입력" in selected_payment:
                payment_method = ""
            else:
                payment_method = PAYMENT_METHODS_PRESET[selected_payment]["desc"]
        st.divider()
        st.markdown("#### 포장 상세")
        pk1, pk2, pk3 = st.columns(3)
        with pk1:
            nw = st.number_input("**순중량(kg)**", value=500.0)
            cbm_doc = st.number_input("**용적(CBM)**", value=2.5)
        with pk2:
            gw = st.number_input("**총중량(KG)**", value=550.0)
            packing_type = st.text_input("**포장형태**", "10 Wooden Pallets")
        with pk3:
            marks = st.text_area("**Shipping Marks**", "N/M\nMADE IN KOREA\nC/NO. 1-10", height=150)
        st.divider()
        submitted = st.form_submit_button("**서류 생성 시작**", use_container_width=True)
    
    with st.expander("**선택사항 설정**", expanded=True):
        opt_col1, opt_col2 = st.columns(2)
        with opt_col1:
            st.markdown("#### 자사 서류 양식 업로드")
            use_custom = st.checkbox("자사 양식 사용하기 (DOCX)", value=False)
            uploaded_templates = None
            if use_custom:
                uploaded_templates = st.file_uploader("Word 파일 업로드 (.docx)", type=['docx'], accept_multiple_files=True, help="파일명에 키워드 포함 시 자동 매칭")
                if uploaded_templates:
                    st.success(f"✅ {len(uploaded_templates)}개 양식 업로드됨")
        with opt_col2:
            st.markdown("#### 직인 첨부하기")
            seal_image = st.file_uploader("회사 직인 이미지", type=["png", "jpg", "jpeg"])
            if seal_image:
                st.image(seal_image, caption="업로드된 직인", width=80)
    
    if submitted:
        if not selected_docs:
            st.warning("⚠️ 생성할 서류를 최소 1개 이상 선택해주세요.")
        else:
            subtotal = qty * unit_price
            ai_advice = ""
            ai_clauses = ""
            if client:
                with st.spinner("🤖 국제무역 전문 AI가 거래를 분석 중..."):
                    try:
                        prompt = f"""당신은 20년 경력의 국제무역 전문 변호사이자 관세사입니다.

[분석 대상 거래 정보]
- 수출자: {shipper}
- 수입자: {consignee}
- 통지처: {notify_party}
- 선적항: {pol}
- 도착항: {pod}
- Incoterms: {incoterms}
- 품목: {item_desc}
- HS Code: {hs_code}
- 수량: {qty} {unit}
- 단가: {selected_currency} {unit_price:,.2f}
- 총액: {selected_currency} {subtotal:,.2f}
- 결제조건: {payment_method}
- 적용 FTA: {selected_fta}
- 적하보험: {insurance_type}
- 순중량/총중량: {nw}kg / {gw}kg
- 용적: {cbm_doc} CBM
- 포장형태: {packing_type}

위 거래 정보를 종합적으로 분석하여 아래 두 가지를 작성해 주십시오.
거래조건을 다시 요약하지 말고, 바로 분석 내용만 작성하세요.

[ADVICE]
한글로 "AI 무역 전략 가이드"를 작성하세요. 번호를 매기고 각 항목에 **굵은 제목**을 붙여주세요.
다음 5가지 관점에서 각각 2~3문장씩 심층 분석해 주세요:

1. **환율 및 관세 리스크**: FTA 적용 여부에 따른 관세율 차이, 환율 변동 리스크 및 헤지 전략
2. **운송 리스크**: 선택된 Incoterms에 따른 위험 이전 시점, 화물 손상/분실 가능성 및 대비책
3. **법적 및 규제 요구사항**: 수입국의 법적 규제, 필요 서류, 통관 시 주의사항
4. **보험 필요성**: 현재 보험 조건의 적정성, 추가 보험 가입 권고 여부
5. **사업 연속성**: 납기 지연, 품질 문제 발생 시 대응 방안 및 위기관리 체계

마지막에 1~2문장으로 결론을 작성하세요.

[CLAUSES]
영문으로 "Special Clauses (특약 조항)"을 작성하세요. 
이 거래에 특화된 법적 효력이 있는 계약 조항 5개를 번호를 매겨 작성하세요.
각 조항은 제목과 상세 내용을 포함해야 합니다:

1. **Price and Payment Terms**: 결제 조건 및 지연 시 페널티
2. **Delivery Terms**: Incoterms 기반 인도 조건 및 위험 이전 시점 명시
3. **Customs and Duties**: 관세 및 통관 책임 소재
4. **Insurance**: 보험 부보 책임 및 범위
5. **Documentation**: 필요 서류 및 제공 의무"""
                        res = client.chat.completions.create(model="gpt-4o", messages=[{"role": "system", "content": "You are an expert international trade lawyer and customs specialist. Provide detailed, practical analysis in Korean for ADVICE section, and professional English contract clauses for CLAUSES section. Do NOT summarize the transaction details - go straight to the analysis."}, {"role": "user", "content": prompt}], max_tokens=2500, temperature=0.4)
                        full_response = res.choices[0].message.content.strip()
                        if "[ADVICE]" in full_response and "[CLAUSES]" in full_response:
                            advice_start = full_response.find("[ADVICE]") + len("[ADVICE]")
                            advice_end = full_response.find("[CLAUSES]")
                            ai_advice = full_response[advice_start:advice_end].strip()
                            clauses_start = full_response.find("[CLAUSES]") + len("[CLAUSES]")
                            ai_clauses = full_response[clauses_start:].strip()
                        else:
                            ai_advice = full_response
                            ai_clauses = full_response
                    except Exception as e:
                        ai_advice = f"AI 분석 오류: {str(e)}\n\n※ OPENAI_API_KEY가 올바르게 설정되었는지 확인하세요."
                        ai_clauses = "Standard Trade Terms Apply per INCOTERMS 2020."
            else:
                ai_advice = "※ AI 분석 기능은 OPENAI_API_KEY 설정 후 사용 가능합니다.\n\n.env 파일에 다음과 같이 추가하세요:\nOPENAI_API_KEY=your-api-key-here"
                ai_clauses = "Standard Trade Terms Apply per INCOTERMS 2020."
            st.session_state['ai_analysis_result'] = ai_advice
            st.session_state['ai_clauses_result'] = ai_clauses
            st.session_state['selected_docs_list'] = selected_docs
            term_code = incoterms.split()[0].upper() 
            departure_terms = ["EXW", "FCA", "FAS", "FOB"]
            if term_code in departure_terms:
                final_incoterms = f"{term_code} {pol}"
            else:
                final_incoterms = f"{term_code} {pod}"
            doc_data = {
                'shipper': shipper, 
                'consignee': consignee, 
                'shipper_addr': shipper.split('\n')[1] if len(shipper.split('\n')) > 1 else '', 
                'consignee_addr': consignee.split('\n')[1] if len(consignee.split('\n')) > 1 else '', 
                'notify_party': notify_party, 
                'inv_no': inv_no, 
                'incoterms': final_incoterms,
                'payment': payment_method, 
                'item_desc': item_desc, 
                'hs_code': hs_code, 
                'qty': qty, 
                'unit_price': unit_price, 
                'unit': unit, 
                'currency': selected_currency, 
                'pol': pol, 
                'pod': pod, 
                'vessel': vessel, 
                'lc_bank': lc_bank, 
                'nw': nw, 
                'gw': gw, 
                'cbm': cbm_doc, 
                'packing_type': packing_type, 
                'marks': marks, 
                'ai_clauses': ai_clauses, 
                'total_amount': subtotal, 
                'delivery_date': (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d'), 
                'shipment_date': (datetime.now() + timedelta(days=45)).strftime('%Y-%m-%d'), 
                'advising_bank': 'Beneficiary Bank', 
                'fta': selected_fta, 
                'insurance': insurance_type
            }
            st.session_state['doc_data'] = doc_data
            seal_path = None
            if seal_image:
                seal_path = io.BytesIO(seal_image.read())
                seal_image.seek(0)
            doc_creators = {"Commercial Invoice": "create_commercial_invoice", "Packing List": "create_packing_list", "Sales Contract": "create_sales_contract", "Proforma Invoice": "create_proforma_invoice", "L/C Application": "create_lc_application", "Purchase Order": "create_purchase_order"}
            with st.spinner("📝 서류 생성 중..."):
                st.session_state['generated_docs'] = {}
                for doc_type in selected_docs:
                    final_doc = None
                    used_method = "표준 양식"
                    if use_custom and uploaded_templates:
                        for tmpl in uploaded_templates:
                            fname = tmpl.name.lower()
                            keywords = {"commercial invoice": ["commercial", "invoice", "ci"], "packing list": ["packing", "pl"], "sales contract": ["sales", "contract", "sc"], "proforma invoice": ["proforma", "pi"], "l/c application": ["lc", "l/c", "credit"], "purchase order": ["purchase", "order", "po"]}
                            for kw in keywords.get(doc_type.lower(), []):
                                if kw in fname:
                                    final_doc = fill_custom_template(tmpl, doc_data, seal_path)
                                    used_method = f"자사 양식 ({tmpl.name})"
                                    break
                            if final_doc:
                                break
                    if final_doc is None:
                        generator = ProfessionalDocGenerator(doc_data, seal_path, trade_mode=trade_mode)
                        final_doc = getattr(generator, doc_creators[doc_type])()
                    doc_bytes = get_document_bytes(final_doc)
                    st.session_state['generated_docs'][doc_type] = {'bytes': doc_bytes, 'method': used_method, 'filename': f"{doc_data['inv_no']}_{doc_type.replace(' ', '_')}.docx"}
                    if seal_path:
                        seal_path.seek(0)
            st.session_state['show_results'] = True
            st.success("✅ 모든 서류가 성공적으로 생성되었습니다!")
    if st.session_state.get('show_results', False) and st.session_state.get('doc_data'):
        st.markdown("---")
        col_ai, col_doc = st.columns(2)
        with col_ai:
            st.markdown("#### 💡 AI 무역 전략 가이드")
            current_ai = st.session_state.get('ai_analysis_result', '')
            current_clauses = st.session_state.get('ai_clauses_result', '')
            st.info(current_ai)
            st.markdown("#### 📜 AI 무역 분석 & 조항")
            st.markdown(f"""<div style="background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%); border-radius: 12px; padding: 20px; border: 1px solid #bae6fd;">{current_clauses.replace(chr(10), '<br>').replace('**', '')}</div>""", unsafe_allow_html=True)
        with col_doc:
            st.markdown("#### 📄 생성된 서류")
            generated_docs = st.session_state.get('generated_docs', {})
            selected_list = st.session_state.get('selected_docs_list', [])
            cols = st.columns(2)
            for i, doc_type in enumerate(selected_list):
                if doc_type in generated_docs:
                    doc_info = generated_docs[doc_type]
                    with cols[i % 2]:
                        st.markdown(f"""<div style="background: #ffffff; padding: 15px; border-radius: 12px; border: 1px solid #e0e4ef; text-align: center; box-shadow: 0 2px 8px rgba(0,0,0,0.05); margin-bottom: 10px;"><b>📄 {doc_type}</b><br><small style="color:#666;">{doc_info['method']}</small></div>""", unsafe_allow_html=True)
                        st.download_button(label=f"⬇️ 다운로드", data=doc_info['bytes'], file_name=doc_info['filename'], mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_{doc_type}_{i}_persistent", use_container_width=True)

st.divider()
st.markdown("""
<div style='text-align: center; color: #718096; font-size: 0.9em;'>
    <p>Global E-commerce All In One Solution</p>
    <p>Developed by Seyeon Global Connect</p>
</div>
""", unsafe_allow_html=True)