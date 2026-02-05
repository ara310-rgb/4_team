import streamlit as st
import os
import time
import requests
from dotenv import load_dotenv
from openai import OpenAI
from pytrends.request import TrendReq
import json
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import html

# --- 페이지 설정 ---
st.set_page_config(
    page_title="Global SEO Marketing Pro", 
    page_icon="🚢", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== 사이드바 네비게이션 ====================
import base64

# ✅ 사이드바 CSS 추가 (기존 CSS와 통합)
st.markdown("""
    <style>
    /* 사이드바 로고 스타일 */
    .logo-box{
        background: rgba(255,255,255,0.6);
        border: 1px solid #e5e7eb;
        border-radius: 14px;
        padding: 14px 12px;
        margin-bottom: 10px;
        text-align:center;
    }
    .logo-img{
        max-width: 150px;
        width: 100%;
        height: auto;
        display:block;
        margin: 0 auto;
    }
    .small-muted{
        color:#64748b;
        font-size: 0.85rem;
        font-weight: 600;
        letter-spacing: 0.2px;
    }
    /* Streamlit 기본 네비게이션 숨김 */
    [data-testid="stSidebarNav"] {
        display: none;
    }
    section[data-testid="stSidebar"] { 
        background: #ffffff !important;
        border-right: 1px solid #e5e7eb;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown("""
    <style>
    /* Streamlit 기본 네비게이션 숨김 */
    [data-testid="stSidebarNav"] {
        display: none;
    }
    
    /* 사이드바 배경 */
    section[data-testid="stSidebar"] { 
        background: #ffffff !important;
        border-right: 1px solid #e5e7eb;
    }
    
    /* 사이드바 버튼 스타일 (✅ hover 효과 추가) */
    .stButton>button {
        background: #051161 !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: 8px;
        padding: 10px 14px;
        font-weight: 700;
        transition: 0.3s;
    }
    .stButton>button:hover {
        background: rgba(5,17,97,0.85) !important;
        box-shadow: 0 4px 12px rgba(5,17,97,0.3) !important;
    }
    
    /* 사이드바 로고 스타일 */
    .logo-box{
        background: rgba(255,255,255,0.6);
        border: 1px solid #e5e7eb;
        border-radius: 14px;
        padding: 14px 12px;
        margin-bottom: 10px;
        text-align:center;
    }
    .logo-img{
        max-width: 150px;
        width: 100%;
        height: auto;
        display:block;
        margin: 0 auto;
    }
    .small-muted{
        color:#64748b;
        font-size: 0.85rem;
        font-weight: 600;
        letter-spacing: 0.2px;
    }
    </style>
""", unsafe_allow_html=True)

with st.sidebar:
    # 1) 국가별 수출입 데이터
    with st.expander("1) 해외진출 전략 허브", expanded=False):
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("시장동향", use_container_width=True, key="nav_cn_1"):
                st.switch_page("pages/macro_1.py")
        with col2:
            if st.button("전략분석", use_container_width=True, key="nav_cn_2"):
                st.switch_page("pages/micro_1.py")
        with col3:
            if st.button("규제진단", use_container_width=True, key="nav_cn_3"):
                st.switch_page("pages/mac_mic_1.py")

    # 2) SEO 서비스
    with st.expander("2) SEO 서비스", expanded=False):
        if st.button("바로가기", use_container_width=True, key="nav_news"):
            st.switch_page("pages/junghyun.py")

    # 3) AI 바이어 매칭 서비스
    with st.expander("3) AI 바이어 매칭 서비스", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            if st.button("바이어 찾기", use_container_width=True, key="nav_ai_1"):
                st.switch_page("pages/03_ai_chatbot.py")
        with col2:
            if st.button("전시회 일정", use_container_width=True, key="nav_ai_2"):
                st.switch_page("pages/buyer_maps.py")

    # 4) 환율 정보 확인
    with st.expander("4) 환율 정보 확인", expanded=False):
        if st.button("바로가기", use_container_width=True, key="nav_ex"):
            st.switch_page("pages/exchange_rate.py")

    # 5) 무역 서류 자동 완성
    with st.expander("5) 무역 서류 자동 완성", expanded=False):
        if st.button("바로가기", use_container_width=True, key="nav_fx"):
            st.switch_page("pages/auto_docs.py")

    # 로고 영역
    logo_path = "assets/logo.png"
    if os.path.exists(logo_path):
        logo_b64 = base64.b64encode(open(logo_path, "rb").read()).decode()
        st.markdown(
            f"""
            <div class="logo-box">
              <img class="logo-img" src="data:image/png;base64,{logo_b64}" alt="logo"/>
              <div class="small-muted" style="margin-top:8px; text-align:center;">
                KITA AX MASTER TEAM4
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("---")

    # ✅ 홈으로 돌아가기 버튼
    if st.button("🏠 홈으로 돌아가기", use_container_width=True, key="go_home_sidebar"):
        st.switch_page("dashboard.py")

# --- CSS 스타일 (향상된 디자인) ---
st.markdown("""
    <style>
    .main { background-color: #f5f7fa; }
    .stButton>button { 
        width: 100%; 
        background: #051161;
        color: white; 
        font-weight: bold; 
        border: none; 
        padding: 14px; 
        border-radius: 10px;
        font-size: 16px;
        transition: all 0.3s;
    }
    .stButton>button:hover { 
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(102, 126, 234, 0.4);
    }
    .keyword-card { 
        background: white; 
        padding: 24px; 
        border-radius: 16px; 
        border: 1px solid #e1e8ed; 
        box-shadow: 0 4px 6px rgba(0,0,0,0.07);
        margin-bottom: 20px;
    }
    .metric-box {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        color: white;
        padding: 15px;
        border-radius: 12px;
        text-align: center;
        margin: 10px 0;
    }
    .keyword-tag {
        display: inline-block;
        background-color: #e3f2fd;
        color: #1976d2;
        padding: 8px 16px;
        border-radius: 20px;
        margin: 5px;
        font-weight: 500;
    }

    .keyword-eng {
        display: block;
        color: #8a8f98;
        font-size: 12px;
        margin-top: 4px;
        font-weight: 400;
    }
    h1, h2, h3, h4 { color: #1a202c; font-weight: 700; }
    .info-badge {
        background-color: #fff3cd;
        color: #856404;
        padding: 4px 12px;
        border-radius: 12px;
        font-size: 0.85em;
        font-weight: 600;
    }
    </style>
    """, unsafe_allow_html=True)

# --- 환경 변수 로드 ---
load_dotenv()
OPENAI_API_KEY = os.getenv("OPEN_API_KEY") or os.getenv("OPENAI_API_KEY")
SERPAPI_KEY = os.getenv("SERPAPI_KEY")
DEEPL_API_KEY = os.getenv("DEEPL_API_KEY")

# ✅ 영어권 국가 정의(한 번만)
ENGLISH_COUNTRIES = ['US', 'GB', 'AU', 'CA', 'NZ', 'SG', 'IE', 'ZA', 'NG', 'PH', 'IN']


# --- 메인 헤더 ---
st.title("🚢 SEO 마케팅 & 콘텐츠 생성")
st.markdown("""
<div style='background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); padding: 20px; border-radius: 12px; color: white; margin-bottom: 30px;'>
    <h4 style='color: white; margin: 0;'> 분석 방식</h4>
    <ul style='margin: 10px 0; padding-left: 20px;'>
        <li>경쟁사 유료 광고 키워드 역분석</li>
        <li>브랜드명·용량 자동 제외한 순수 마케팅 키워드 추출</li>
        <li>STP 분석 및 핵심 타겟 소비층 설정</li>
        <li>아마존·자사몰·SNS 콘텐츠 현지화 생성</li>
    </ul>
</div>
""", unsafe_allow_html=True)

# --- 입력 섹션 ---
st.subheader("제품 정보 입력")

col1, col2, col3 = st.columns([1.2, 1.5, 1])

with col1:
    input_type = st.selectbox(
        "입력 방식 선택",
        ["제품 키워드", "HS Code"],
        help="HS Code는 6자리 국제 관세 코드입니다"
    )

with col2:
    if input_type == "HS Code":
        placeholder_text = "예: 190230 (파스타), 851830 (헤드폰)"
    else:
        placeholder_text = "예: 무선 이어폰, 프로틴 파우더"
    
    user_input = st.text_input(
        "제품명 또는 HS Code",
        placeholder=placeholder_text
    )

with col3:
    # 국가를 지역별로 그룹화
    country_options = {
        "북미": ["US", "CA", "MX"],
        "유럽": ["GB", "DE", "FR", "IT", "ES", "NL", "BE", "AT", "CH", "SE", "NO", "DK", "FI", "PL", "IE", "PT", "GR", "CZ", "RO", "HU"],
        "아시아": ["JP", "KR", "CN", "SG", "TW", "HK", "TH", "MY", "ID", "VN", "PH", "IN"],
        "오세아니아": ["AU", "NZ"],
        "중동": ["AE", "SA", "IL", "TR"],
        "남미": ["BR", "AR", "CL", "CO"],
        "아프리카": ["ZA", "EG", "NG"]
    }
    
    # 전체 국가 리스트 (알파벳 순)
    all_countries = []
    for countries in country_options.values():
        all_countries.extend(countries)
    all_countries = sorted(list(set(all_countries)))
    
    target_country = st.selectbox(
        "타겟 국가",
        all_countries,
        help="마케팅 대상 국가를 선택하세요 (60개국 이상 지원)"
    )

analyze_btn = st.button("전체 시장 분석 시작", use_container_width=True)

# --- 유틸리티 함수 ---
def get_language_code(country_code):
    """국가별 언어 코드 매핑 (확장)"""
    lang_map = {
        # 북미
        'US': 'en', 'CA': 'en', 'MX': 'es',
        # 유럽
        'GB': 'en', 'IE': 'en',
        'DE': 'de', 'AT': 'de', 'CH': 'de',
        'FR': 'fr', 'BE': 'fr',
        'IT': 'it', 'ES': 'es', 'PT': 'pt',
        'NL': 'nl', 'SE': 'sv', 'NO': 'no', 'DK': 'da', 'FI': 'fi',
        'PL': 'pl', 'CZ': 'cs', 'RO': 'ro', 'HU': 'hu', 'GR': 'el',
        # 아시아
        'JP': 'ja', 'KR': 'ko', 'CN': 'zh-CN', 'TW': 'zh-TW', 'HK': 'zh-HK',
        'SG': 'en', 'TH': 'th', 'MY': 'ms', 'ID': 'id', 'VN': 'vi', 'PH': 'en',
        'IN': 'en',
        # 오세아니아
        'AU': 'en', 'NZ': 'en',
        # 중동
        'AE': 'ar', 'SA': 'ar', 'IL': 'he', 'TR': 'tr',
        # 남미
        'BR': 'pt', 'AR': 'es', 'CL': 'es', 'CO': 'es',
        # 아프리카
        'ZA': 'en', 'EG': 'ar', 'NG': 'en'
    }
    return lang_map.get(country_code, 'en')

def get_language_name(country_code):
    """국가별 언어 이름 매핑 (확장)"""
    lang_map = {
        # 북미
        'US': 'English (US)', 'CA': 'English (CA)', 'MX': 'Spanish (Mexico)',
        # 유럽
        'GB': 'English (UK)', 'IE': 'English (Ireland)',
        'DE': 'German', 'AT': 'German (Austria)', 'CH': 'German (Switzerland)',
        'FR': 'French', 'BE': 'French (Belgium)',
        'IT': 'Italian', 'ES': 'Spanish', 'PT': 'Portuguese',
        'NL': 'Dutch', 'SE': 'Swedish', 'NO': 'Norwegian', 'DK': 'Danish', 'FI': 'Finnish',
        'PL': 'Polish', 'CZ': 'Czech', 'RO': 'Romanian', 'HU': 'Hungarian', 'GR': 'Greek',
        # 아시아
        'JP': 'Japanese', 'KR': 'Korean', 'CN': 'Chinese (Simplified)', 
        'TW': 'Chinese (Traditional)', 'HK': 'Chinese (Hong Kong)',
        'SG': 'English (Singapore)', 'TH': 'Thai', 'MY': 'Malay', 
        'ID': 'Indonesian', 'VN': 'Vietnamese', 'PH': 'English (Philippines)',
        'IN': 'English (India)',
        # 오세아니아
        'AU': 'English (Australia)', 'NZ': 'English (New Zealand)',
        # 중동
        'AE': 'Arabic (UAE)', 'SA': 'Arabic (Saudi Arabia)', 'IL': 'Hebrew', 'TR': 'Turkish',
        # 남미
        'BR': 'Portuguese (Brazil)', 'AR': 'Spanish (Argentina)', 
        'CL': 'Spanish (Chile)', 'CO': 'Spanish (Colombia)',
        # 아프리카
        'ZA': 'English (South Africa)', 'EG': 'Arabic (Egypt)', 'NG': 'English (Nigeria)'
    }
    return lang_map.get(country_code, 'English')

def get_country_name(country_code):
    """국가 코드 → 국가명 (확장)"""
    country_map = {
        # 북미
        'US': 'United States', 'CA': 'Canada', 'MX': 'Mexico',
        # 유럽
        'GB': 'United Kingdom', 'IE': 'Ireland',
        'DE': 'Germany', 'AT': 'Austria', 'CH': 'Switzerland',
        'FR': 'France', 'BE': 'Belgium',
        'IT': 'Italy', 'ES': 'Spain', 'PT': 'Portugal',
        'NL': 'Netherlands', 'SE': 'Sweden', 'NO': 'Norway', 'DK': 'Denmark', 'FI': 'Finland',
        'PL': 'Poland', 'CZ': 'Czech Republic', 'RO': 'Romania', 'HU': 'Hungary', 'GR': 'Greece',
        # 아시아
        'JP': 'Japan', 'KR': 'South Korea', 'CN': 'China', 'TW': 'Taiwan', 'HK': 'Hong Kong',
        'SG': 'Singapore', 'TH': 'Thailand', 'MY': 'Malaysia', 
        'ID': 'Indonesia', 'VN': 'Vietnam', 'PH': 'Philippines', 'IN': 'India',
        # 오세아니아
        'AU': 'Australia', 'NZ': 'New Zealand',
        # 중동
        'AE': 'United Arab Emirates', 'SA': 'Saudi Arabia', 'IL': 'Israel', 'TR': 'Turkey',
        # 남미
        'BR': 'Brazil', 'AR': 'Argentina', 'CL': 'Chile', 'CO': 'Colombia',
        # 아프리카
        'ZA': 'South Africa', 'EG': 'Egypt', 'NG': 'Nigeria'
    }
    return country_map.get(country_code, country_code)

# --- 핵심 로직 함수 ---

def get_seed_keyword(client, user_input, input_type, target_country_code):
    """
    HS Code 또는 제품명 → 현지 언어 표준 품목명 변환
    관세사(Customs Broker) 페르소나 사용
    """
    lang_name = get_language_name(target_country_code)
    
    system_prompt = """You are an expert International Trade Specialist and Customs Broker with deep knowledge of HS Code classifications and global product nomenclature."""
    
    if input_type == "HS Code":
        user_prompt = f"""
Analyze HS CODE: '{user_input}'
Target Market: {target_country_code} ({lang_name})

TASK: Identify the **OFFICIAL GENERIC PRODUCT CATEGORY NAME** for this HS Code.

⛔ STRICT PROHIBITIONS:
- NO specific brand names (e.g., 'Sony', 'Nike', 'Barilla')
- NO niche flavors or variants unless they define the main category
- NO product model numbers

✅ REQUIRED OUTPUT:
- Most common, generic search term consumers use in {lang_name}
- Must be the broadest commercially viable category

Examples:
- Input '190230' → Output: 'Pasta' (or 'パスタ' for JP)
- Input '851830' → Output: 'Headphones' (or 'ヘッドフォン' for JP)
- Input '330210' → Output: 'Perfume'

Output Format (2 words only, separated by comma):
NativeLanguageKeyword, EnglishKeyword
"""
    else:
        user_prompt = f"""
Product Name: '{user_input}'
Target Market: {target_country_code} ({lang_name})

TASK: Translate this product into the most common generic search term in {lang_name}.

⛔ Remove any brand names or specific measurements.
✅ Output the broadest product category term.

Output Format (2 words only, separated by comma):
NativeLanguageKeyword, EnglishKeyword
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3
        )
        content = response.choices[0].message.content.strip()
        content = content.replace('*', '').replace('"', '').replace("'", '').replace('`', '')
        
        if ',' in content:
            parts = [x.strip() for x in content.split(',', 1)]
            return parts[0], parts[1] if len(parts) > 1 else parts[0]
        
        return content.strip(), content.strip()
    except Exception as e:
        st.error(f"키워드 식별 실패: {str(e)}")
        return None, None


def fetch_comprehensive_serpapi_data(search_term, country_code):
    """
    SerpApi 통합 데이터 수집:
    1. Google Shopping (경쟁사 제품명)
    2. Organic Search (상위 랭크 사이트의 메타 설명)
    3. Ads (유료 광고 키워드 - 전환 키워드의 보고)
    4. People Also Ask (관련 질문 - 롱테일 키워드)
    5. Related Searches (구글 추천 키워드)
    """
    if not SERPAPI_KEY:
        return None
    
    lang_code = get_language_code(country_code)
    collected_data = {
        "shopping_titles": [],
        "ads_headlines": [],
        "organic_snippets": [],
        "people_also_ask": [],
        "related_searches": []
    }
    
    # 1. Google Shopping 검색
    try:
        params_shopping = {
            "engine": "google_shopping",
            "q": search_term,
            "gl": country_code,
            "hl": lang_code,
            "api_key": SERPAPI_KEY,
            "num": 30
        }
        
        res = requests.get("https://serpapi.com/search", params=params_shopping, timeout=15).json()
        
        if "shopping_results" in res:
            for item in res["shopping_results"][:30]:
                title = item.get("title", "")
                if title:
                    collected_data["shopping_titles"].append(title)
        
        time.sleep(0.5)  # Rate limiting
    except Exception as e:
        st.warning(f"Shopping 데이터 수집 실패: {str(e)}")
    
    # 2. Google Organic + Ads 검색
    try:
        params_organic = {
            "engine": "google",
            "q": search_term,
            "gl": country_code,
            "hl": lang_code,
            "api_key": SERPAPI_KEY,
            "num": 20
        }
        
        res = requests.get("https://serpapi.com/search", params=params_organic, timeout=15).json()
        
        # 2-1. 유료 광고 (Ads) - 전환 키워드의 핵심
        if "ads" in res:
            for ad in res["ads"][:10]:
                headline = ad.get("title", "") or ad.get("headline", "")
                if headline:
                    collected_data["ads_headlines"].append(headline)
        
        # 2-2. Organic 검색 결과 (메타 설명)
        if "organic_results" in res:
            for item in res["organic_results"][:10]:
                snippet = item.get("snippet", "")
                if snippet:
                    collected_data["organic_snippets"].append(snippet)
        
        # 2-3. People Also Ask (질문 형태 키워드)
        if "related_questions" in res:
            for q in res["related_questions"][:10]:
                question = q.get("question", "")
                if question:
                    collected_data["people_also_ask"].append(question)
        
        # 2-4. Related Searches (구글 추천 검색어)
        if "related_searches" in res:
            for rs in res["related_searches"][:10]:
                query = rs.get("query", "")
                if query:
                    collected_data["related_searches"].append(query)
        
    except Exception as e:
        st.warning(f"Organic/Ads 데이터 수집 실패: {str(e)}")
    
    return collected_data



def extract_high_intent_keywords(client, serp_data, country, product_name):
    """
    SerpApi 데이터에서 고의도 마케팅 키워드(영어 원본) 추출
    - 브랜드/리테일러/용량/정보성 키워드 제외
    - 반환: (comma-separated English string, is_fallback)
    """
    is_fallback = False

    if serp_data and any(serp_data.values()):
        context_parts = []
        if serp_data.get("shopping_titles"):
            context_parts.append(f"Shopping Titles: {serp_data['shopping_titles'][:20]}")
        if serp_data.get("ads_headlines"):
            context_parts.append(f"Paid Ads (High-Intent): {serp_data['ads_headlines'][:10]}")
        if serp_data.get("related_searches"):
            context_parts.append(f"Related Searches: {serp_data['related_searches']}")
        context = "\n".join(context_parts)
        task_desc = "Extract the top 10 high-volume commercial keywords from the above data."
    else:
        is_fallback = True
        context = "No real-time SerpApi data available."
        task_desc = f"Generate top 10 commercial keywords for '{product_name}' based on general market knowledge for {country}."

    prompt = f"""
Role: E-commerce SEO Keyword Analyst for {get_country_name(country)}.

Context:
{context}

Task: {task_desc}

⛔ CRITICAL EXCLUSION RULES (MUST FOLLOW):
1. NO BRAND NAMES (Nike, Adidas, Samsung, Apple, etc.)
2. NO RETAILER/PLATFORM NAMES (Amazon, Walmart, Target, etc.)
3. NO SPECIFIC MEASUREMENTS (500g, 200ml, 12 pack, etc.)
4. NO INFORMATIONAL TERMS (recipe, how to, tutorial, history, calories, nutrition facts)

✅ OUTPUT REQUIREMENTS:
- Output EXACTLY 10 keywords
- Output ENGLISH ONLY
- Each keyword: 1~5 words, commercial intent
- Each keyword on a NEW LINE
- No numbering, no bullet symbols, no extra text
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4
        )
        raw = (response.choices[0].message.content or "").strip()

        # 1) 안전한 후처리: 프롬프트/헤더가 섞여 나오면 제거
        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        bad_prefix = ("role:", "context:", "task:", "output", "data available:")
        cleaned = []
        for ln in lines:
            lnl = ln.lower()
            if any(lnl.startswith(p) for p in bad_prefix):
                continue
            ln = re.sub(r'^\d+[\.\)]\s*', '', ln)  # 1. 제거
            ln = ln.strip("•-–— \t")
            ln = ln.rstrip(",")
            if ln:
                cleaned.append(ln)

        # 2) 줄이 한 줄로 붙어서 왔으면 콤마 기준도 분해
        if len(cleaned) <= 1:
            parts = re.split(r'[,\n]+', raw)
            cleaned = [p.strip().rstrip(",") for p in parts if p.strip()]

        # 3) 길이/중복 정리
        seen = set()
        final = []
        for kw in cleaned:
            norm = re.sub(r'\s+', ' ', kw).strip()
            if not norm:
                continue
            key = norm.lower()
            if key in seen:
                continue
            seen.add(key)
            final.append(norm)
            if len(final) >= 10:
                break

        if not final:
            return "분석 실패", True

        return ", ".join(final), is_fallback

    except Exception as e:
        st.error(f"키워드 추출 실패: {str(e)}")
        return "분석 실패", True

def extract_longtail_keywords(client, serp_data, country, product_name):
    """
    People Also Ask 및 Related Searches에서 롱테일 키워드(영어) 추출
    - 반환: List[str] (최대 8개)
    """
    if not serp_data:
        return []

    paa_data = serp_data.get("people_also_ask", [])
    related_data = serp_data.get("related_searches", [])

    if not paa_data and not related_data:
        return []

    prompt = f"""
Role: SEO Long-tail Keyword Specialist for {get_country_name(country)}.

Data:
- People Also Ask: {paa_data}
- Related Searches: {related_data}

Product (English): {product_name}

TASK:
Generate EXACTLY 8 long-tail keywords (3-6 words) in ENGLISH.
They should include:
- informational intent
- comparison intent
- purchase intent

⛔ Exclude brand names, retailer/platform names, and specific measurements.
✅ Output rules:
- ENGLISH ONLY
- Each keyword on a NEW LINE
- No numbering, no bullets, no extra text
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4
        )
        raw = (response.choices[0].message.content or "").strip()

        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        bad_prefix = ("role:", "data:", "task:", "output", "product")
        cleaned = []
        for ln in lines:
            lnl = ln.lower()
            if any(lnl.startswith(p) for p in bad_prefix):
                continue
            ln = re.sub(r'^\d+[\.\)]\s*', '', ln)
            ln = ln.strip("•-–— \t")
            ln = ln.rstrip(",")
            if ln:
                cleaned.append(ln)

        # 콤마로 한 줄에 뭉쳐온 경우 분해
        if len(cleaned) <= 1:
            parts = re.split(r'[,\n]+', raw)
            cleaned = [p.strip().rstrip(",") for p in parts if p.strip()]

        # 중복 제거 + 8개 제한
        seen = set()
        final = []
        for kw in cleaned:
            norm = re.sub(r'\s+', ' ', kw).strip()
            if not norm:
                continue
            key = norm.lower()
            if key in seen:
                continue
            seen.add(key)
            final.append(norm)
            if len(final) >= 8:
                break

        return final

    except Exception:
        return []
def fetch_google_trends(seed_keyword, geo_code):
    """Google Trends 데이터 수집"""
    pytrends = TrendReq(hl='en-US', tz=360)
    
    for timeframe in ['today 12-m', 'today 3-m', 'today 1-m']:
        try:
            pytrends.build_payload([seed_keyword], timeframe=timeframe, geo=geo_code)
            related = pytrends.related_queries()
            
            if related and seed_keyword in related:
                top_df = related[seed_keyword]['top']
                if top_df is not None and not top_df.empty:
                    return top_df.head(10)['query'].tolist()
            
            time.sleep(1)
        except:
            continue
    
    return []


# ============================================
# DeepL 번역 함수
# ============================================

def translate_with_deepl(text, target_country):
    """
    DeepL API를 사용한 고품질 번역
    target_country: 국가 코드 (예: 'JP', 'KR', 'FR')
    """
    if not DEEPL_API_KEY:
        return None
    
    # 국가 코드 → DeepL 언어 코드 매핑
    deepl_lang_map = {
        'JP': 'JA', 'KR': 'KO', 'CN': 'ZH', 'TW': 'ZH', 'HK': 'ZH',
        'FR': 'FR', 'DE': 'DE', 'ES': 'ES', 'IT': 'IT',
        'PT': 'PT-PT', 'BR': 'PT-BR',
        'NL': 'NL', 'PL': 'PL', 'RU': 'RU', 'TR': 'TR',
        'SE': 'SV', 'NO': 'NB', 'DK': 'DA', 'FI': 'FI',
        'GR': 'EL', 'CZ': 'CS', 'RO': 'RO', 'HU': 'HU',
        'ID': 'ID', 'AR': 'AR', 'TH': 'TH', 'VI': 'VI'
    }
    
    target_lang = deepl_lang_map.get(target_country)
    if not target_lang:
        return None  # DeepL이 지원하지 않는 언어
    
    api_url = "https://api-free.deepl.com/v2/translate"
    
    try:
        response = requests.post(
            api_url,
            data={
                'auth_key': DEEPL_API_KEY,
                'text': text,
                'target_lang': target_lang,
                'source_lang': 'EN'  # 소스는 영어로 고정
            },
            timeout=15
        )
        
        if response.status_code == 200:
            result = response.json()
            return result['translations'][0]['text']
        else:
            return None
    except:
        return None

def translate_with_gpt(client, text, target_country):
    """
    GPT를 사용한 번역 (DeepL 실패 시 폴백)
    """
    lang_name = get_language_name(target_country)

    prompt = f"""Translate the following English text to {lang_name}.
Keep the marketing tone and style.
Do NOT add any explanations, just provide the translated text.

Text to translate:
{text}"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        return response.choices[0].message.content.strip()
    except:
        return None



def translate_to_english_with_deepl(text, source_country):
    """
    DeepL로 (현지어 -> 영어) 번역. 지원하지 않으면 None.
    - DeepL은 언어 자동감지가 가능하므로 source_lang을 지정하지 않음
    """
    if not DEEPL_API_KEY:
        return None

    # DeepL이 지원하는 타겟 국가만 빠르게 체크 (기존 맵 재사용)
    deepl_supported = {
        'JP','KR','CN','TW','HK','FR','DE','ES','IT','PT','BR','NL','PL','RU','TR',
        'SE','NO','DK','FI','GR','CZ','RO','HU','ID','AR','TH','VI'
    }
    if source_country not in deepl_supported:
        return None

    api_url = "https://api-free.deepl.com/v2/translate"
    try:
        response = requests.post(
            api_url,
            data={
                'auth_key': DEEPL_API_KEY,
                'text': text,
                'target_lang': 'EN'  # 영어로
            },
            timeout=15
        )
        if response.status_code == 200:
            result = response.json()
            return result['translations'][0]['text']
        return None
    except:
        return None


def translate_to_english(client, text, source_country):
    """
    (현지어 -> 영어) 번역. DeepL 우선, 실패 시 GPT 폴백.
    - 출력은 '영어만' 나오도록 강제
    """
    if not text:
        return ""

    # 1) DeepL 우선
    en = translate_to_english_with_deepl(text, source_country)
    if en:
        return en.strip()

    # 2) GPT 폴백
    lang_name = get_language_name(source_country)
    prompt = f"""Translate the following text from {lang_name} to English.
Rules:
- Output ENGLISH ONLY.
- Keep the original meaning and marketing tone.
- Do NOT add explanations, labels, or extra commentary.

Text:
{text}"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        en = (response.choices[0].message.content or "").strip()

        # 영어만 나오도록 1회 후처리: 비라틴 문자가 많이 섞이면 재요청
        non_latin = sum(1 for ch in en if ord(ch) > 127)
        if len(en) > 0 and (non_latin / max(len(en), 1)) > 0.15:
            retry_prompt = f"""Rewrite the following into ENGLISH ONLY.
Do not include any non-English words.

Text:
{en}"""
            response2 = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": retry_prompt}],
                temperature=0.1
            )
            en = (response2.choices[0].message.content or "").strip()

        return en
    except:
        return None


def clean_trends_keywords(client, keywords, product_en, country):
    """
    Google Trends 결과에서 브랜드/리테일러/용량 등을 제거하고 10개로 정리.
    - 입력 언어를 유지(현지어면 현지어로 유지)
    """
    if not keywords:
        return []

    prompt = f"""
You are a strict SEO keyword cleaner for {get_country_name(country)}.

Given this list of search queries:
{keywords}

Remove:
- brand names (any company/brand)
- retailer/platform names (Amazon, etc.)
- specific measurements (ml, g, oz, pack, etc.)
- overly specific model numbers

Keep:
- generic, commercial, high-intent queries

Return up to 10 cleaned queries.
OUTPUT RULES:
- Return ONLY the cleaned queries
- One per line
- Keep the SAME LANGUAGE as the input list
"""
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2
        )
        raw = (resp.choices[0].message.content or "").strip()
        lines = [ln.strip().strip("•-–— \t").rstrip(",") for ln in raw.splitlines() if ln.strip()]
        # 헤더 제거
        lines = [ln for ln in lines if not ln.lower().startswith(("output", "return", "cleaned"))]
        # 중복 제거
        seen = set()
        final = []
        for ln in lines:
            key = ln.lower()
            if key in seen:
                continue
            seen.add(key)
            final.append(ln)
            if len(final) >= 10:
                break
        return final if final else keywords[:10]
    except:
        return keywords[:10]


# ============================================
# 타겟 소비층 분석 (Segmentation)
# ============================================

def generate_target_audience_analysis(client, keywords, product, country, serp_data):
    """
    타겟 소비층 분석 (Demographics & Persona Summary 위주로 압축)
    """
    country_name = get_country_name(country)
    
    # 컨텍스트 준비
    context_parts = [f"Product: {product}", f"High-Intent Keywords: {keywords}"]
    if serp_data and serp_data.get("ads_headlines"):
        context_parts.append(f"Competitor Ads: {serp_data['ads_headlines'][:3]}")
    
    context = "\n".join(context_parts)
    
    prompt = f"""
You are a Market Segmentation Analyst for {country_name}.

CONTEXT:
{context}

TASK: Provide a CONCISE Demographic & Persona analysis. Do NOT write long paragraphs. Use bullet points.

OUTPUT FORMAT:

### Core Demographics
* **Age Group:** [e.g., 25-34]
* **Gender:** [e.g., Female skewing]
* **Income Level:** [e.g., Upper-middle class]
* **Location:** [e.g., Urban areas]
* **Occupation:** [Key job titles]

### Psychographics & Pain Points
* **Core Values:** [3 key values]
* **Main Pain Point:** [The #1 problem they want to solve]
* **Buying Motivation:** [Why they buy this specific product]

### Quick Marketing Insight
* [One sentence on the best angle to approach this audience]

Write in English. Keep it brief and actionable.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.6
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"타겟 분석 실패: {str(e)}"

def generate_high_quality_content(client, target_analysis, keywords, product, country):
    """
    타겟 소비층 기반 고품질 마케팅 콘텐츠 생성 (항상 영어로 생성)
    """
    country_name = get_country_name(country)
    
    prompt = f"""
You are an Elite E-commerce Copywriter for the {country_name} market.
You specialize in high-conversion sales copy that drives immediate purchases.

TARGET AUDIENCE INSIGHTS:
{target_analysis}

Product: {product}
Keywords to Integrate: {keywords}

=== CRITICAL LANGUAGE INSTRUCTION ===
- Write ENTIRELY in ENGLISH (it will be translated later if needed)
- Use culturally appropriate expressions for {country_name} market
- All keywords, descriptions, and hashtags MUST be in ENGLISH

TASK: Write 3 types of marketing assets with a "Premium Brand" tone.

=== 1. AMAZON BULLET POINTS (5 Bullets) ===
Style Guide:
- **Format:** [BENEFIT HEADER IN CAPS] - [Detailed explanation]
- **Header:** 2-4 words, punchy, benefit-driven (e.g., "INSTANT PAIN RELIEF", "ULTRA-DURABLE MATERIAL")
- **Body:** 2-3 sentences. Explain WHY this matters. Use sensory words.
- **Goal:** Overcome objections and create desire.
- **Length:** 200-250 characters per bullet (optimal for mobile indexing).

=== 2. D2C PRODUCT DESCRIPTION (Storytelling) ===
Style Guide:
- **Tone:** Empathetic, authoritative, yet accessible.
- **Structure:**
  1. **Headline:** A promise of transformation (e.g., "Reclaim Your Comfort Today")
  2. **The Problem:** Empathize with the user's struggle (1 paragraph)
  3. **The Solution:** Introduce the product as the hero (1 paragraph)
  4. **The "Why Us":** Highlight unique selling points/technology (1 paragraph)
  5. **Closing:** A strong call to action.

=== 3. SOCIAL MEDIA POST (Instagram/Facebook) ===
Style Guide:
- **Hook:** A question or bold statement to stop the scroll.
- **Vibe:** Lifestyle-focused, not just product specs.
- **Structure:** Hook -> Relatable Scenario -> Product Benefit -> CTA.
- **Emojis:** Use relevant emojis to break up text (but don't overdo it).
- **Hashtags:** Mix of broad (#Marketing) and niche (#EcoFriendlyLiving) tags.

=== OUTPUT FORMAT (Strictly follow this) ===

---AMAZON_START---
- [HEADER]: [Content]
- [HEADER]: [Content]
- [HEADER]: [Content]
- [HEADER]: [Content]
- [HEADER]: [Content]
---AMAZON_END---

---D2C_START---
# [Headline]

[Body Paragraph 1 - Problem]

[Body Paragraph 2 - Solution]

[Body Paragraph 3 - Features]

[CTA]
---D2C_END---

---SOCIAL_START---
[Hook]

[Body]

[CTA]

[Hashtags]
---SOCIAL_END---

IMPORTANT: Write EVERYTHING in ENGLISH.
Ensure keywords are woven in naturally for SEO.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.75
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"콘텐츠 생성 실패: {str(e)}"


def parse_persona_content(content_text):
    """
    페르소나 기반 콘텐츠를 섹션별로 파싱
    Returns: (amazon_text, d2c_text, social_text)
    """
    import re
    
    # 기본값
    amazon_text = ""
    d2c_text = ""
    social_text = ""
    
    try:
        # Amazon 섹션 추출
        amazon_match = re.search(r'---AMAZON_START---(.*?)---AMAZON_END---', content_text, re.DOTALL)
        if amazon_match:
            amazon_text = amazon_match.group(1).strip()
        
        # D2C 섹션 추출
        d2c_match = re.search(r'---D2C_START---(.*?)---D2C_END---', content_text, re.DOTALL)
        if d2c_match:
            d2c_text = d2c_match.group(1).strip()
        
        # Social 섹션 추출
        social_match = re.search(r'---SOCIAL_START---(.*?)---SOCIAL_END---', content_text, re.DOTALL)
        if social_match:
            social_text = social_match.group(1).strip()
        
        # 마커가 없는 경우 (fallback) - # 기준으로 파싱
        if not amazon_text or not d2c_text or not social_text:
            sections = content_text.split('#')
            
            for section in sections:
                section_lower = section.lower()
                
                if 'amazon' in section_lower and not amazon_text:
                    # "Amazon Bullet Points" 헤더 제거
                    lines = section.split('\n')
                    content_lines = [line for line in lines if 'amazon' not in line.lower() and line.strip()]
                    amazon_text = '\n'.join(content_lines).strip()
                
                elif ('d2c' in section_lower or 'product description' in section_lower) and not d2c_text:
                    lines = section.split('\n')
                    content_lines = [line for line in lines if 'd2c' not in line.lower() and 'product description' not in line.lower() and line.strip()]
                    d2c_text = '\n'.join(content_lines).strip()
                
                elif ('social' in section_lower or 'instagram' in section_lower or 'facebook' in section_lower) and not social_text:
                    lines = section.split('\n')
                    content_lines = [line for line in lines if 'social' not in line.lower() and 'instagram' not in line.lower() and 'facebook' not in line.lower() and line.strip()]
                    social_text = '\n'.join(content_lines).strip()
        
        # 여전히 비어있으면 전체 텍스트를 각각에 할당 (최후의 수단)
        if not amazon_text:
            amazon_text = "콘텐츠 파싱 실패 - 원본 전체를 표시합니다."
        if not d2c_text:
            d2c_text = "콘텐츠 파싱 실패 - 원본 전체를 표시합니다."
        if not social_text:
            social_text = "콘텐츠 파싱 실패 - 원본 전체를 표시합니다."
        
    except Exception as e:
        # 파싱 실패 시 전체 텍스트 반환
        error_msg = f"파싱 오류: {str(e)}"
        return error_msg, error_msg, error_msg
    
    return amazon_text, d2c_text, social_text


# ============================================
# 메인 실행 로직
# ============================================

# 1. 세션 스테이트 초기화 (데이터 유지용)
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None

# ✅ 입력값 변경 시 편집 위젯(text_area) 상태가 이전 값을 끌고 오는 문제 방지
if 'last_run_signature' not in st.session_state:
    st.session_state.last_run_signature = None

def reset_output_widgets():
    # Streamlit text_area는 key가 같으면 이전 값(session_state)을 계속 유지하므로 삭제 필요
    for k in ("amazon_edit_area", "d2c_edit_area", "sns_edit_area"):
        if k in st.session_state:
            del st.session_state[k]

# ✅ 입력(제품/국가/입력방식) 변경 시 이전 분석 결과/편집상태 자동 초기화 (브라우저 새로고침 효과)
current_input_signature = f"{input_type}|{user_input}|{target_country}"
if 'last_input_signature' not in st.session_state:
    st.session_state.last_input_signature = current_input_signature
elif st.session_state.last_input_signature != current_input_signature:
    # 사용자가 제품/국가 등을 바꾼 상태에서 이전 결과가 화면에 남아 혼동되는 것을 방지
    reset_output_widgets()
    st.session_state.analysis_result = None
    st.session_state.last_input_signature = current_input_signature

# 2. 분석 버튼 실행 로직
if analyze_btn:
    # ★★★ 새 분석 시작 시 이전 세션/위젯 초기화 ★★★
    reset_output_widgets()
    st.session_state.last_run_signature = f"{input_type}|{user_input}|{target_country}"
    st.session_state.analysis_result = None
    
    if not user_input or not target_country:
        st.warning("⚠️ 제품명/HS Code와 타겟 국가를 모두 입력해주세요.")
    elif not OPENAI_API_KEY:
        st.error("❌ OpenAI API Key가 설정되지 않았습니다. `.env` 파일을 확인하세요.")
    else:
        # OpenAI 클라이언트 초기화
        client = OpenAI(api_key=OPENAI_API_KEY)
        
        # 프로그레스 바
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # --- [Step 1~4: 데이터 수집 및 분석] ---
        
        status_text.text("🔍 Step 1/7: 제품명 식별 중...")
        progress_bar.progress(10)
        native_kw, english_kw = get_seed_keyword(client, user_input, input_type, target_country)
        
        if not native_kw:
            st.error("식별 실패")
            st.stop()

        # 언어 보정
        if target_country in ['US', 'GB', 'AU', 'CA', 'NZ', 'SG', 'IE', 'ZA', 'NG', 'PH', 'IN'] and any(ord(c) > 127 for c in native_kw):
            native_kw = english_kw

        status_text.text("🌐 Step 2/7: 시장 데이터 수집 중 ...")
        progress_bar.progress(30)
        serp_data = fetch_comprehensive_serpapi_data(native_kw, target_country)

        status_text.text("💎 Step 3/7: 고의도 키워드 분석 중...")
        high_intent_kw, is_fallback = extract_high_intent_keywords(client, serp_data, target_country, english_kw)
        longtail_kw = extract_longtail_keywords(client, serp_data, target_country, english_kw)
        progress_bar.progress(45)

        # 키워드 현지화 처리
        is_english_country = target_country in ENGLISH_COUNTRIES

        # (영어 원본을 별도로 보관) - 비영어권에서 '현지어 + 영어' 함께 표시하기 위함
        high_intent_kw_en_list = []
        high_intent_kw_local_list = []
        longtail_kw_en = list(longtail_kw) if longtail_kw else []
        longtail_kw_local = list(longtail_kw) if longtail_kw else []

        if high_intent_kw and "분석 실패" not in high_intent_kw:
            high_intent_kw_en_list = [k.strip() for k in high_intent_kw.split(',') if k.strip()]

        if not is_english_country:
            status_text.text("🌐 Step 3/7: 키워드 현지화 중...")

            # High-Intent 키워드: '영어 리스트'를 기준으로 개별 번역하여 정렬/의미 보존
            if high_intent_kw_en_list:
                for kw in high_intent_kw_en_list:
                    translated = translate_with_deepl(kw, target_country)
                    if not translated:
                        translated = translate_with_gpt(client, kw, target_country)
                    high_intent_kw_local_list.append(translated if translated else kw)
                    time.sleep(0.1)

                # 기존 호환용(문자열)도 유지
                high_intent_kw = ", ".join(high_intent_kw_local_list)

            # Long-tail 키워드: 영어 리스트를 기준으로 개별 번역
            if longtail_kw_en:
                translated_longtail = []
                for kw in longtail_kw_en:
                    translated = translate_with_deepl(kw, target_country)
                    if not translated:
                        translated = translate_with_gpt(client, kw, target_country)
                    translated_longtail.append(translated if translated else kw)
                    time.sleep(0.1)
                longtail_kw_local = translated_longtail
                longtail_kw = longtail_kw_local
        else:
            # 영어권: 현지어=영어
            high_intent_kw_local_list = list(high_intent_kw_en_list)
            longtail_kw_local = list(longtail_kw_en)

        progress_bar.progress(50)

        status_text.text("📈 Step 4/7: Google Trends 분석 중...")
        trends_kw_raw = fetch_google_trends(native_kw, target_country)
        # ✅ 브랜드/리테일러/용량 제거 (Trends는 원천 데이터라 필터가 필요)
        trends_kw = clean_trends_keywords(client, trends_kw_raw, english_kw, target_country)
        progress_bar.progress(60)

        # Trends 키워드 영어 의미(비영어권 대시보드 표시용)
        trends_kw_en = []
        if target_country not in ENGLISH_COUNTRIES and trends_kw:
            for kw in trends_kw:
                en = translate_to_english(client, kw, target_country)
                trends_kw_en.append(en if en else kw)
                time.sleep(0.1)
        else:
            trends_kw_en = list(trends_kw) if trends_kw else []

        status_text.text("👥 Step 5/7: 타겟 소비층 분석 중...")
        # ✅ 분석/콘텐츠 생성용 키워드는 항상 영어 원본 사용 (혼용/번역 오류 방지)
        high_intent_kw_for_generation = ", ".join(high_intent_kw_en_list) if high_intent_kw_en_list else english_kw
        target_analysis = generate_target_audience_analysis(client, high_intent_kw_for_generation, english_kw, target_country, serp_data)
        progress_bar.progress(70)

        status_text.text("✍️ Step 6/7: 고품질 마케팅 콘텐츠 생성 중...")
        marketing_content_en = generate_high_quality_content(client, target_analysis, high_intent_kw_for_generation, english_kw, target_country)
        amazon_en, d2c_en, social_en = parse_persona_content(marketing_content_en)
        progress_bar.progress(80)

        status_text.text("🌐 Step 7/7: 현지화 번역 중...")
        
        # 영어권 여부 확인
        is_english_country = target_country in ['US', 'GB', 'AU', 'CA', 'NZ', 'SG', 'IE', 'ZA', 'NG', 'PH', 'IN']
        
        if is_english_country:
            # 영어권 - 번역 불필요
            amazon_final = amazon_en
            d2c_final = d2c_en
            social_final = social_en
            translation_status = "영어 원본"
        else:
            # 비영어권 - DeepL 우선, 실패 시 GPT
            translation_results = []
            
            # 아마존 번역
            amazon_final = translate_with_deepl(amazon_en, target_country)
            if amazon_final:
                translation_results.append("DeepL")
            else:
                amazon_final = translate_with_gpt(client, amazon_en, target_country)
                if amazon_final:
                    translation_results.append("GPT")
                else:
                    amazon_final = amazon_en
                    translation_results.append("실패")
            
            # 자사몰 번역
            d2c_final = translate_with_deepl(d2c_en, target_country)
            if d2c_final:
                translation_results.append("DeepL")
            else:
                d2c_final = translate_with_gpt(client, d2c_en, target_country)
                if d2c_final:
                    translation_results.append("GPT")
                else:
                    d2c_final = d2c_en
                    translation_results.append("실패")
            
            # SNS 번역
            social_final = translate_with_deepl(social_en, target_country)
            if social_final:
                translation_results.append("DeepL")
            else:
                social_final = translate_with_gpt(client, social_en, target_country)
                if social_final:
                    translation_results.append("GPT")
                else:
                    social_final = social_en
                    translation_results.append("실패")
            
            # 번역 상태 결정 (워드 파일용만)
            deepl_count = translation_results.count("DeepL")
            gpt_count = translation_results.count("GPT")
            
            if deepl_count == 3:
                translation_status = "DeepL"
            elif gpt_count == 3:
                translation_status = "GPT"
            elif deepl_count + gpt_count == 3:
                translation_status = "DeepL + GPT"
            else:
                translation_status = "부분 번역"

        # ✅ 비영어권: 현지어 최종본의 '영어 의미'를 별도로 생성(혼용/깨짐 방지)
        amazon_meaning_en = None
        d2c_meaning_en = None
        social_meaning_en = None
        if not is_english_country:
            amazon_meaning_en = translate_to_english(client, amazon_final, target_country) or amazon_en
            d2c_meaning_en = translate_to_english(client, d2c_final, target_country) or d2c_en
            social_meaning_en = translate_to_english(client, social_final, target_country) or social_en

        progress_bar.progress(100)
        time.sleep(0.5)
        status_text.empty()
        progress_bar.empty()

        # 결과 세션에 저장 (화면 리프레시 되더라도 유지)
        st.session_state.analysis_result = {
            'native_kw': native_kw,
            'english_kw': english_kw,
            'target_country': target_country,
            'high_intent_kw': high_intent_kw,
            'high_intent_kw_en_list': high_intent_kw_en_list,
            'high_intent_kw_local_list': high_intent_kw_local_list,
            'longtail_kw': longtail_kw,
            'longtail_kw_en': longtail_kw_en,
            'longtail_kw_local': longtail_kw_local,
            'trends_kw': trends_kw,
            'trends_kw_en': trends_kw_en,
            'target_analysis': target_analysis,
            'amazon_en': amazon_en,
            'd2c_en': d2c_en,
            'social_en': social_en,
            'amazon_meaning_en': amazon_meaning_en,
            'd2c_meaning_en': d2c_meaning_en,
            'social_meaning_en': social_meaning_en,
            'amazon_final': amazon_final,
            'd2c_final': d2c_final,
            'social_final': social_final,
            'translation_status': translation_status,
            'is_fallback': is_fallback,
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
        }

# 3. 결과 화면 출력 (세션에 데이터가 있을 경우 항상 표시)
if st.session_state.analysis_result:
    data = st.session_state.analysis_result
    saved_country = data['target_country']
    
    # ★★★ 번역 상태 메시지 제거 - 성공 메시지만 표시 ★★★
    st.success(f"✅ **{get_country_name(saved_country)}** 시장 분석 완료: **{data['english_kw']}** ({data['native_kw']})")

    st.divider()

    # --- 타겟 소비층 분석 섹션 ---
    st.header("타겟 소비층 분석 (Demographics)")
    with st.expander("📊 핵심 인구통계 및 인사이트 보기", expanded=True):
        st.markdown(data['target_analysis'])

    st.divider()

    # --- 키워드 섹션 ---
    st.header("추출된 키워드")
    k_tab1, k_tab2, k_tab3 = st.tabs(["💎 High-Intent", "🔍 Long-tail", "📈 Trends"])
    
    with k_tab1:
        st.subheader("고의도 마케팅 키워드")
        st.caption("브랜드명·용량 제외, 전환율 높은 키워드")
        
        if "분석 실패" not in data['high_intent_kw']:
            # ★★★ 아랍어/히브리어 등 RTL 언어 처리 개선 ★★★
            is_english_country_ui = saved_country in ENGLISH_COUNTRIES
            
            if (not is_english_country_ui) and data.get('high_intent_kw_local_list') and data.get('high_intent_kw_en_list'):
                # 현지어 + 영어(회색) 함께 표시
                pairs = list(zip(data['high_intent_kw_local_list'], data['high_intent_kw_en_list']))
                tags_html = "".join([
                    f'<span class="keyword-tag">🔥 {html.escape(local)}<span class="keyword-eng">{html.escape(en)}</span></span>'
                    for local, en in pairs
                ])
                st.markdown(tags_html, unsafe_allow_html=True)
            else:
                keywords_list = [k.strip() for k in data['high_intent_kw'].split(',') if k.strip()]
                tags_html = "".join([f'<span class="keyword-tag">🔥 {html.escape(kw)}</span>' for kw in keywords_list])
                st.markdown(tags_html, unsafe_allow_html=True)
        else:
            st.error("키워드 추출 실패")

    with k_tab2:
        st.subheader("🔍 롱테일 키워드 (SEO 최적화)")
        st.caption("'People Also Ask' 및 '관련 검색어' 기반")
        
        if data['longtail_kw']:
            is_english_country_ui = saved_country in ENGLISH_COUNTRIES
            
            if (not is_english_country_ui) and data.get('longtail_kw_local') and data.get('longtail_kw_en'):
                for idx, (local_kw, en_kw) in enumerate(zip(data['longtail_kw_local'], data['longtail_kw_en']), 1):
                    st.markdown(f"**{idx}.** {local_kw}")
                    st.markdown(f"<div style='color:#8a8f98;font-size:12px;margin-left:22px;'>{html.escape(en_kw)}</div>", unsafe_allow_html=True)
            else:
                for idx, kw in enumerate(data['longtail_kw'], 1):
                    st.markdown(f"**{idx}.** {kw}")
        else:
            st.warning("데이터 부족")

    with k_tab3:
        st.subheader("📈 Google Trends 인기 검색어")
        st.caption("실제 소비자 검색 행동 데이터")
        
        if data['trends_kw']:
            is_english_country_ui = saved_country in ENGLISH_COUNTRIES
            
            if (not is_english_country_ui) and data.get('trends_kw_en') and data.get('trends_kw'):
                for idx, (local_kw, en_kw) in enumerate(zip(data['trends_kw'], data['trends_kw_en']), 1):
                    st.markdown(f"**{idx}.** {local_kw}")
                    st.markdown(f"<div style='color:#8a8f98;font-size:12px;margin-left:22px;'>{html.escape(en_kw)}</div>", unsafe_allow_html=True)
            else:
                for idx, kw in enumerate(data['trends_kw'], 1):
                    st.markdown(f"**{idx}.** {kw}")
        else:
            st.warning("데이터 부족")

    st.divider()

    # --- 마케팅 콘텐츠 섹션 ---
    st.header("마케팅 콘텐츠 생성")

    content_tab1, content_tab2, content_tab3 = st.tabs([
        "🛒 아마존 불렛포인트",
        "🌐 자사몰 상세페이지",
        "📱 SNS 마케팅 피드"
    ])

    with content_tab1:
        with st.container(border=True):
            st.subheader("아마존 제품 설명")
            edited_amazon = st.text_area(
                "아마존 불렛포인트 수정",
                value=data['amazon_final'],
                height=400,
                key="amazon_edit_area",
                label_visibility="collapsed"
            )

            is_english_country_ui = saved_country in ENGLISH_COUNTRIES
            if (not is_english_country_ui) and data.get('amazon_meaning_en'):
                st.markdown(
                    f"<div style='color:#8a8f98;font-size:12px; margin-top:10px; white-space:pre-wrap;'><b>English meaning</b><br>{html.escape(data['amazon_meaning_en'])}</div>",
                    unsafe_allow_html=True
                )

    with content_tab2:
        with st.container(border=True):
            st.subheader("🌐 자사몰 제품 설명")
            edited_d2c = st.text_area(
                "자사몰 설명 수정",
                value=data['d2c_final'],
                height=500,
                key="d2c_edit_area",
                label_visibility="collapsed"
            )

            is_english_country_ui = saved_country in ENGLISH_COUNTRIES
            if (not is_english_country_ui) and data.get('d2c_meaning_en'):
                st.markdown(
                    f"<div style='color:#8a8f98;font-size:12px; margin-top:10px; white-space:pre-wrap;'><b>English meaning</b><br>{html.escape(data['d2c_meaning_en'])}</div>",
                    unsafe_allow_html=True
                )

    with content_tab3:
        with st.container(border=True):
            st.subheader("SNS 마케팅 포스트")
            edited_sns = st.text_area(
                "SNS 피드 수정",
                value=data['social_final'],
                height=400,
                key="sns_edit_area",
                label_visibility="collapsed"
            )

            is_english_country_ui = saved_country in ENGLISH_COUNTRIES
            if (not is_english_country_ui) and data.get('social_meaning_en'):
                st.markdown(
                    f"<div style='color:#8a8f98;font-size:12px; margin-top:10px; white-space:pre-wrap;'><b>English meaning</b><br>{html.escape(data['social_meaning_en'])}</div>",
                    unsafe_allow_html=True
                )

    st.divider()

    # --- 다운로드 섹션 ---
    st.subheader("결과 다운로드 ")

    # 엑셀 파일 생성
    def create_excel_report_dynamic(amazon_txt, d2c_txt, sns_txt):
        output = BytesIO()
        is_english_country_dl = saved_country in ENGLISH_COUNTRIES

        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # 1. 요약 (워드 파일의 '제품 정보'와 동일하게 확장)
            summary_items = [
                ('제품명 (현지어)', data.get('native_kw', '')),
                ('제품명 (영어)', data.get('english_kw', '')),
                ('타겟 국가', get_country_name(saved_country)),
                ('언어', get_language_name(saved_country)),
                ('번역 방식', data.get('translation_status', '')),
                ('생성 일시', data.get('timestamp', ''))
            ]
            pd.DataFrame({
                '항목': [x[0] for x in summary_items],
                '내용': [x[1] for x in summary_items]
            }).to_excel(writer, sheet_name='📋 요약', index=False)

            # 2. 타겟 분석
            pd.DataFrame({'내용': [data.get('target_analysis', '')]}).to_excel(writer, sheet_name='🎯 타겟분석', index=False)

            # 3. 고의도 키워드 (+ 영어 의미: 비영어권만)
            if (not is_english_country_dl) and data.get('high_intent_kw_local_list') and data.get('high_intent_kw_en_list'):
                df_kw = pd.DataFrame({
                    '현지어 키워드': data['high_intent_kw_local_list'],
                    '영어 의미': data['high_intent_kw_en_list']
                })
            else:
                kw_list = [k.strip() for k in (data.get('high_intent_kw', '') or '').split(',') if k.strip()]
                df_kw = pd.DataFrame({'키워드': kw_list})
            df_kw.to_excel(writer, sheet_name='💎 키워드', index=False)

            # 4. 콘텐츠 (+ 영어 의미: 비영어권만)
            if not is_english_country_dl:
                pd.DataFrame({
                    '현지어(최종)': [amazon_txt],
                    '영어 의미': [data.get('amazon_meaning_en') or data.get('amazon_en') or '']
                }).to_excel(writer, sheet_name='🛒 아마존', index=False)

                pd.DataFrame({
                    '현지어(최종)': [d2c_txt],
                    '영어 의미': [data.get('d2c_meaning_en') or data.get('d2c_en') or '']
                }).to_excel(writer, sheet_name='🌐 자사몰', index=False)

                pd.DataFrame({
                    '현지어(최종)': [sns_txt],
                    '영어 의미': [data.get('social_meaning_en') or data.get('social_en') or '']
                }).to_excel(writer, sheet_name='📱 SNS', index=False)
            else:
                pd.DataFrame({'아마존 불렛': [amazon_txt]}).to_excel(writer, sheet_name='🛒 아마존', index=False)
                pd.DataFrame({'자사몰 상세': [d2c_txt]}).to_excel(writer, sheet_name='🌐 자사몰', index=False)
                pd.DataFrame({'SNS 피드': [sns_txt]}).to_excel(writer, sheet_name='📱 SNS', index=False)

        output.seek(0)
        return output


    # 워드 파일 생성
    def create_word_report_dynamic(amazon_txt, d2c_txt, sns_txt):
        doc = Document()
        is_english_country_dl = saved_country in ENGLISH_COUNTRIES

        def clean_and_add_text(document, text_content):
            if not text_content:
                return
            lines = text_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                clean_line = line.replace('### ', '').replace('## ', '').replace('**', '').replace('__', '')
                clean_line = re.sub(r'[^\w\s\uAC00-\uD7A3\u0600-\u06FF.,!?%&()\-:;\'\"]', '', clean_line).strip()
                if not clean_line:
                    continue
                if line.startswith('* ') or line.startswith('- ') or line.startswith('• '):
                    clean_text = clean_line.lstrip('*-• ').strip()
                    document.add_paragraph(clean_text, style='List Bullet')
                else:
                    document.add_paragraph(clean_line)

        title = doc.add_heading('SEO 마케팅 분석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('제품 정보', 1)
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_data = [
            ('제품명 (현지어)', data.get('native_kw', '')),
            ('제품명 (영어)', data.get('english_kw', '')),
            ('타겟 국가', get_country_name(saved_country)),
            ('언어', get_language_name(saved_country)),
            ('번역 방식', data.get('translation_status', '')),
            ('생성 일시', data.get('timestamp', ''))
        ]

        for idx, (label, value) in enumerate(info_data):
            info_table.rows[idx].cells[0].text = label
            info_table.rows[idx].cells[1].text = str(value)

        doc.add_paragraph()
        doc.add_heading('타겟 소비층 분석 (Segmentation)', 1)
        clean_and_add_text(doc, data.get('target_analysis', ''))
        doc.add_page_break()

        doc.add_heading('고의도 마케팅 키워드', 1)
        doc.add_paragraph('브랜드명·용량 제외, 전환율 높은 키워드').italic = True

        if "분석 실패" not in (data.get('high_intent_kw', '') or ''):
            if (not is_english_country_dl) and data.get('high_intent_kw_local_list') and data.get('high_intent_kw_en_list'):
                pairs = list(zip(data['high_intent_kw_local_list'], data['high_intent_kw_en_list']))
                for local_kw, en_kw in pairs:
                    doc.add_paragraph(str(local_kw), style='List Number')
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.35)
                    run = p.add_run(str(en_kw))
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(138, 143, 152)
            else:
                keywords_list = [k.strip() for k in (data.get('high_intent_kw', '') or '').split(',') if k.strip()]
                for kw in keywords_list:
                    clean_kw = re.sub(r'[^\w\s\uAC00-\uD7A3\u0600-\u06FF.,!?%&()\-:;\'\"]', '', kw).replace('**', '')
                    doc.add_paragraph(clean_kw, style='List Number')
        else:
            doc.add_paragraph('키워드 추출 실패')

        doc.add_paragraph()
        doc.add_heading('롱테일 키워드 (SEO 최적화)', 1)
        doc.add_paragraph('People Also Ask 및 관련 검색어 기반').italic = True
        if data.get('longtail_kw_local'):
            if (not is_english_country_dl) and data.get('longtail_kw_en'):
                for local_kw, en_kw in zip(data['longtail_kw_local'], data['longtail_kw_en']):
                    doc.add_paragraph(str(local_kw), style='List Number')
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.35)
                    run = p.add_run(str(en_kw))
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(138, 143, 152)
            else:
                for kw in data['longtail_kw_local']:
                    clean_kw = re.sub(r'[^\w\s\uAC00-\uD7A3\u0600-\u06FF.,!?%&()\-:;\'\"]', '', str(kw)).replace('**', '')
                    doc.add_paragraph(clean_kw, style='List Number')
        else:
            doc.add_paragraph('데이터 없음')

        doc.add_paragraph()
        doc.add_heading('Google Trends 인기 검색어', 1)
        doc.add_paragraph('실제 소비자 검색 행동 데이터').italic = True
        if data.get('trends_kw'):
            if (not is_english_country_dl) and data.get('trends_kw_en'):
                for local_kw, en_kw in zip(data['trends_kw'], data['trends_kw_en']):
                    doc.add_paragraph(str(local_kw), style='List Number')
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.35)
                    run = p.add_run(str(en_kw))
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(138, 143, 152)
            else:
                for kw in data['trends_kw']:
                    clean_kw = re.sub(r'[^\w\s\uAC00-\uD7A3\u0600-\u06FF.,!?%&()\-:;\'\"]', '', str(kw)).replace('**', '')
                    doc.add_paragraph(clean_kw, style='List Number')
        else:
            doc.add_paragraph('데이터 수집 실패')

        doc.add_page_break()
        doc.add_heading('마케팅 콘텐츠 (최종)', 1)

        # 아마존
        doc.add_heading('아마존 블랙보드 불렛포인트', 2)
        clean_and_add_text(doc, amazon_txt)
        if not is_english_country_dl:
            doc.add_paragraph()
            doc.add_heading('English Meaning', 3)
            clean_and_add_text(doc, data.get('amazon_meaning_en') or data.get('amazon_en') or '')
            st.space("small")

        # 자사몰
        doc.add_paragraph()
        doc.add_heading('자사몰 제품 Description', 2)
        clean_and_add_text(doc, d2c_txt)
        if not is_english_country_dl:
            doc.add_paragraph()
            doc.add_heading('English Meaning', 3)
            clean_and_add_text(doc, data.get('d2c_meaning_en') or data.get('d2c_en') or '')
            st.space("small")

        # SNS
        doc.add_paragraph()
        doc.add_heading('SNS 마케팅 피드', 2)
        clean_and_add_text(doc, sns_txt)
        if not is_english_country_dl:
            doc.add_paragraph()
            doc.add_heading('English Meaning', 3)
            clean_and_add_text(doc, data.get('social_meaning_en') or data.get('social_en') or '')
            st.space("small")

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output


    # 다운로드 버튼
    col_dl1, col_dl2 = st.columns(2)
    
    excel_file = create_excel_report_dynamic(edited_amazon, edited_d2c, edited_sns)
    col_dl1.download_button(
        label="📊 엑셀 파일 다운로드",
        data=excel_file,
        file_name=f"Marketing_Report_{data['english_kw']}_{saved_country}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )
    
    word_file = create_word_report_dynamic(edited_amazon, edited_d2c, edited_sns)
    col_dl2.download_button(
        label="📄 워드 파일 다운로드",
        data=word_file,
        file_name=f"Marketing_Report_{data['english_kw']}_{saved_country}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

# --- Footer ---
st.divider()
st.markdown("""
<div style='text-align: center; color: #718096; font-size: 0.9em;'>
    <p>Global E-commerce All In One Solution</p>
    <p>Developed by Seyeon Global Connect</p>
</div>
""", unsafe_allow_html=True)